"""
–£–ª—É—á—à–µ–Ω–Ω—ã–π –º–æ–¥—É–ª—å –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤ .docx –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ö
–†–∞–±–æ—Ç–∞–µ—Ç –Ω–∞–ø—Ä—è–º—É—é —Å ZIP —Å—Ç—Ä—É–∫—Ç—É—Ä–æ–π –¥–ª—è –Ω–∞–¥–µ–∂–Ω–æ–≥–æ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
"""

import os
import io
import tempfile
import logging
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class ImageInfo:
    """–ö–ª–∞—Å—Å –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –æ–± –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–∏"""
    image_id: str
    image_data: bytes
    image_format: str
    width: Optional[float] = None
    height: Optional[float] = None
    paragraph_index: Optional[int] = None
    rel_id: Optional[str] = None
    filename: Optional[str] = None


@dataclass
class ImageElement:
    """–ö–ª–∞—Å—Å –¥–ª—è —Ö—Ä–∞–Ω–µ–Ω–∏—è —ç–ª–µ–º–µ–Ω—Ç–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è (—Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç—å —Å —É–¥–∞–ª–µ–Ω–Ω—ã–º image_processor)"""
    image_id: str
    image_data: bytes
    image_format: str
    width: Optional[int] = None
    height: Optional[int] = None
    paragraph_index: Optional[int] = None
    is_inline: bool = True
    description: Optional[str] = None
    alt_text: Optional[str] = None


class ImprovedImageProcessor:
    """–£–ª—É—á—à–µ–Ω–Ω—ã–π –∫–ª–∞—Å—Å –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ö"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.temp_dir = None
        self.images: List[ImageInfo] = []
        
    def extract_images_from_docx(self, docx_path: str) -> List[ImageInfo]:
        """
        –ò–∑–≤–ª–µ–∫–∞–µ—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –Ω–∞–ø—Ä—è–º—É—é –∏–∑ ZIP —Å—Ç—Ä—É–∫—Ç—É—Ä—ã .docx —Ñ–∞–π–ª–∞
        
        Args:
            docx_path: –ü—É—Ç—å –∫ .docx —Ñ–∞–π–ª—É
            
        Returns:
            –°–ø–∏—Å–æ–∫ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –æ–± –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è—Ö
        """
        images = []
        
        try:
            # –°–æ–∑–¥–∞–µ–º –≤—Ä–µ–º–µ–Ω–Ω—É—é –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—é
            if not self.temp_dir:
                self.temp_dir = tempfile.mkdtemp(prefix='docx_images_')
            
            # –û—Ç–∫—Ä—ã–≤–∞–µ–º .docx –∫–∞–∫ ZIP –∞—Ä—Ö–∏–≤
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # –ü–æ–ª—É—á–∞–µ–º —Å–ø–∏—Å–æ–∫ –≤—Å–µ—Ö —Ñ–∞–π–ª–æ–≤ –≤ –∞—Ä—Ö–∏–≤–µ
                file_list = docx_zip.namelist()
                
                # –ò—â–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –ø–∞–ø–∫–µ word/media/
                media_files = [f for f in file_list if f.startswith('word/media/')]
                
                # –ü–∞—Ä—Å–∏–º relationships –¥–ª—è –ø–æ–ª—É—á–µ–Ω–∏—è —Å–≤—è–∑–µ–π
                relationships = self._parse_relationships(docx_zip)
                
                # –ü–∞—Ä—Å–∏–º –æ—Å–Ω–æ–≤–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç –¥–ª—è –ø–æ–∏—Å–∫–∞ –ø–æ–∑–∏—Ü–∏–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
                image_positions = self._parse_document_for_images(docx_zip)
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –¥–ª—è –ª–æ–≥–∏—Ä–æ–≤–∞–Ω–∏—è
                self._last_relationships = relationships
                self._last_positions = image_positions
                
                self.logger.info(f"–ù–∞–π–¥–µ–Ω–æ {len(media_files)} –º–µ–¥–∏–∞ —Ñ–∞–π–ª–æ–≤ –∏ {len(relationships)} relationships")
                self.logger.info(f"–ù–∞–π–¥–µ–Ω–æ {len(image_positions)} –ø–æ–∑–∏—Ü–∏–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤ –¥–æ–∫—É–º–µ–Ω—Ç–µ")
                
                # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∫–∞–∂–¥–æ–µ –º–µ–¥–∏–∞ —Ñ–∞–π–ª
                positioned_images = 0
                unpositioned_images = 0
                
                for media_file in media_files:
                    try:
                        # –ò–∑–≤–ª–µ–∫–∞–µ–º –¥–∞–Ω–Ω—ã–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                        image_data = docx_zip.read(media_file)
                        
                        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ñ–æ—Ä–º–∞—Ç
                        image_format = self._detect_image_format(image_data)
                        
                        if image_format == 'unknown':
                            self.logger.warning(f"–ù–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {media_file}")
                            continue
                        
                        # –°–æ–∑–¥–∞–µ–º ID
                        image_id = f"extracted_{len(images) + 1}"
                        filename = os.path.basename(media_file)
                        
                        # –ò—â–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –ø–æ–∑–∏—Ü–∏–∏
                        rel_id = self._find_rel_id_for_media(media_file, relationships)
                        paragraph_index = image_positions.get(rel_id)
                        
                        # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è —Å –ø–æ–∑–∏—Ü–∏—è–º–∏
                        if paragraph_index is not None:
                            positioned_images += 1
                            self.logger.info(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {filename} -> rel_id: {rel_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {paragraph_index}")
                        else:
                            unpositioned_images += 1
                            self.logger.warning(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {filename} -> rel_id: {rel_id} -> –ø–æ–∑–∏—Ü–∏—è –Ω–µ –Ω–∞–π–¥–µ–Ω–∞")
                        
                        # –ü–æ–ª—É—á–∞–µ–º —Ä–∞–∑–º–µ—Ä—ã –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                        width, height = self._get_image_dimensions(image_data)
                        
                        # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –≤–æ –≤—Ä–µ–º–µ–Ω–Ω—É—é –ø–∞–ø–∫—É
                        temp_path = os.path.join(self.temp_dir, f"{image_id}.{image_format}")
                        with open(temp_path, 'wb') as f:
                            f.write(image_data)
                        
                        # –°–æ–∑–¥–∞–µ–º –æ–±—ä–µ–∫—Ç –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –æ–± –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–∏
                        image_info = ImageInfo(
                            image_id=image_id,
                            image_data=image_data,
                            image_format=image_format,
                            width=width,
                            height=height,
                            paragraph_index=paragraph_index,
                            rel_id=rel_id,
                            filename=filename
                        )
                        
                        images.append(image_info)
                        self.logger.info(f"–ò–∑–≤–ª–µ—á–µ–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {filename} ({image_format}, {len(image_data)} –±–∞–π—Ç)")
                        
                    except Exception as e:
                        self.logger.warning(f"–û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –º–µ–¥–∏–∞ —Ñ–∞–π–ª–∞ {media_file}: {e}")
                        continue
            
            self.images = images
            self.logger.info(f"–í—Å–µ–≥–æ –∏–∑–≤–ª–µ—á–µ–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {len(images)}")
            self.logger.info(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π —Å –Ω–∞–π–¥–µ–Ω–Ω—ã–º–∏ –ø–æ–∑–∏—Ü–∏—è–º–∏: {positioned_images}")
            self.logger.info(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –±–µ–∑ –ø–æ–∑–∏—Ü–∏–π: {unpositioned_images}")
            
        except Exception as e:
            self.logger.error(f"–û—à–∏–±–∫–∞ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –∏–∑ {docx_path}: {e}")
            
        return images
    
    def _parse_relationships(self, docx_zip: zipfile.ZipFile) -> Dict[str, str]:
        """–ü–∞—Ä—Å–∏—Ç —Ñ–∞–π–ª relationships –¥–ª—è –ø–æ–ª—É—á–µ–Ω–∏—è —Å–≤—è–∑–µ–π –º–µ–∂–¥—É ID –∏ —Ñ–∞–π–ª–∞–º–∏"""
        relationships = {}
        
        try:
            # –ß–∏—Ç–∞–µ–º word/_rels/document.xml.rels
            rels_content = docx_zip.read('word/_rels/document.xml.rels')
            root = ET.fromstring(rels_content)
            
            # –ü–∞—Ä—Å–∏–º relationships
            for rel in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rel_id = rel.get('Id')
                target = rel.get('Target')
                rel_type = rel.get('Type')
                
                # –ò–Ω—Ç–µ—Ä–µ—Å—É—é—Ç —Ç–æ–ª—å–∫–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                if rel_type and 'image' in rel_type.lower():
                    relationships[rel_id] = target
                    
        except Exception as e:
            self.logger.warning(f"–û—à–∏–±–∫–∞ –ø–∞—Ä—Å–∏–Ω–≥–∞ relationships: {e}")
            
        return relationships
    
    def _parse_document_for_images(self, docx_zip: zipfile.ZipFile) -> Dict[str, int]:
        """
        –£–õ–£–ß–®–ï–ù–ù–´–ô –ø–∞—Ä—Å–∏–Ω–≥ –æ—Å–Ω–æ–≤–Ω–æ–≥–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –¥–ª—è –ø–æ–∏—Å–∫–∞ –ø–æ–∑–∏—Ü–∏–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
        –û–±–µ—Å–ø–µ—á–∏–≤–∞–µ—Ç –±–æ–ª–µ–µ —Ç–æ—á–Ω–æ–µ –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –ø–æ–∑–∏—Ü–∏–π –∏ –ª—É—á—à—É—é —Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç—å —Å python-docx
        """
        image_positions = {}
        
        try:
            # –ß–∏—Ç–∞–µ–º word/document.xml
            doc_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(doc_content)
            
            # –ò—â–µ–º –¢–û–õ–¨–ö–û –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã –≤ –æ—Å–Ω–æ–≤–Ω–æ–º —Ç–µ–ª–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞ (–∏—Å–∫–ª—é—á–∞–µ–º headers, footers, etc.)
            body = root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
            if body is None:
                self.logger.warning("–ù–µ –Ω–∞–π–¥–µ–Ω body —ç–ª–µ–º–µ–Ω—Ç –≤ –¥–æ–∫—É–º–µ–Ω—Ç–µ")
                return image_positions
            
            # –ü–æ–ª—É—á–∞–µ–º –í–°–ï –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã –∏–∑ body (–≤–∫–ª—é—á–∞—è –ø—É—Å—Ç—ã–µ)
            # –≠—Ç–æ –≤–∞–∂–Ω–æ –¥–ª—è –ø—Ä–∞–≤–∏–ª—å–Ω–æ–≥–æ —Å–æ–ø–æ—Å—Ç–∞–≤–ª–µ–Ω–∏—è —Å python-docx
            all_paragraphs = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            
            self.logger.info(f"üîç XML –ø–∞—Ä—Å–µ—Ä: –Ω–∞–π–¥–µ–Ω–æ {len(all_paragraphs)} –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –≤ –æ—Å–Ω–æ–≤–Ω–æ–º —Ç–µ–ª–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞")
            print(f"üîç XML –ø–∞—Ä—Å–µ—Ä: –Ω–∞–π–¥–µ–Ω–æ {len(all_paragraphs)} –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –≤ –æ—Å–Ω–æ–≤–Ω–æ–º —Ç–µ–ª–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞")
            
            # –°–Ω–∞—á–∞–ª–∞ –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º –≤—Å–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è –∫–∞—Ä—Ç—ã —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏–π
            paragraph_mapping = []
            significant_para_count = 0
            
            for xml_idx, paragraph in enumerate(all_paragraphs):
                # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞
                text_nodes = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                image_nodes = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                
                # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º —Ç–µ–∫—Å—Ç–æ–≤–æ–µ —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ
                text_content = ''.join(node.text or '' for node in text_nodes).strip()
                has_meaningful_text = len(text_content) > 0
                has_images = len(image_nodes) > 0
                
                # –û–ø—Ä–µ–¥–µ–ª—è–µ–º, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ –ø–∞—Ä–∞–≥—Ä–∞—Ñ –∑–Ω–∞—á–∏–º—ã–º
                is_significant = has_meaningful_text or has_images
                
                if is_significant:
                    paragraph_mapping.append({
                        'xml_index': xml_idx,
                        'docx_index': significant_para_count,
                        'has_text': has_meaningful_text,
                        'has_images': has_images,
                        'text_preview': text_content[:50] + '...' if len(text_content) > 50 else text_content
                    })
                    significant_para_count += 1
                    
                    self.logger.debug(f"–ü–∞—Ä–∞–≥—Ä–∞—Ñ XML[{xml_idx}] -> DOCX[{significant_para_count-1}]: "
                                    f"text={has_meaningful_text}, images={has_images}, "
                                    f"preview='{text_content[:30]}...' if text_content else 'empty'")
            
            self.logger.info(f"üîç –°–æ–∑–¥–∞–Ω–æ —Å–æ–ø–æ—Å—Ç–∞–≤–ª–µ–Ω–∏–µ: {len(paragraph_mapping)} –∑–Ω–∞—á–∏–º—ã—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏–∑ {len(all_paragraphs)} XML –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤")
            print(f"üîç –°–æ–∑–¥–∞–Ω–æ —Å–æ–ø–æ—Å—Ç–∞–≤–ª–µ–Ω–∏–µ: {len(paragraph_mapping)} –∑–Ω–∞—á–∏–º—ã—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏–∑ {len(all_paragraphs)} XML –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤")
            
            # –¢–µ–ø–µ—Ä—å –∏—â–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –∫–∞–∂–¥–æ–º –ø–∞—Ä–∞–≥—Ä–∞—Ñ–µ
            for xml_idx, paragraph in enumerate(all_paragraphs):
                images_in_paragraph = []
                
                # –ù–∞—Ö–æ–¥–∏–º —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–∏–π docx –∏–Ω–¥–µ–∫—Å
                docx_idx = None
                for mapping in paragraph_mapping:
                    if mapping['xml_index'] == xml_idx:
                        docx_idx = mapping['docx_index']
                        break
                
                # –ï—Å–ª–∏ –ø–∞—Ä–∞–≥—Ä–∞—Ñ –Ω–µ –∑–Ω–∞—á–∏–º—ã–π, –ø—Ä–æ–ø—É—Å–∫–∞–µ–º –ø–æ–∏—Å–∫ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
                if docx_idx is None:
                    continue
                
                # === –£–õ–£–ß–®–ï–ù–ù–´–ô –ü–û–ò–°–ö –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô ===
                
                # 1. –ò—â–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ drawing —ç–ª–µ–º–µ–Ω—Ç–∞—Ö (—Å–æ–≤—Ä–µ–º–µ–Ω–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç)
                drawings = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                for drawing in drawings:
                    # –ò—â–µ–º –≤—Å–µ –≤–æ–∑–º–æ–∂–Ω—ã–µ —Ç–∏–ø—ã –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤ drawing
                    blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id:
                            image_positions[embed_id] = docx_idx
                            images_in_paragraph.append(f"drawing:{embed_id}")
                            self.logger.debug(f"–ù–∞–π–¥–µ–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –≤ drawing: {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {docx_idx}")
                
                # 2. –ò—â–µ–º inline –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è (–≤—Å—Ç—Ä–æ–µ–Ω–Ω—ã–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è)  
                inline_shapes = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object')
                for inline_shape in inline_shapes:
                    blips = inline_shape.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id:
                            image_positions[embed_id] = docx_idx
                            images_in_paragraph.append(f"inline:{embed_id}")
                            self.logger.debug(f"–ù–∞–π–¥–µ–Ω–æ inline –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {docx_idx}")
                
                # 3. –ò—â–µ–º pict —ç–ª–µ–º–µ–Ω—Ç—ã (—Å—Ç–∞—Ä—ã–π —Ñ–æ—Ä–º–∞—Ç Word)
                picts = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                for pict in picts:
                    # –ò—â–µ–º —Ä–∞–∑–ª–∏—á–Ω—ã–µ —Ç–∏–ø—ã –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –≤ pict
                    shapes = pict.findall('.//*[@r:id]', namespaces={'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
                    for shape in shapes:
                        embed_id = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if embed_id:
                            image_positions[embed_id] = docx_idx
                            images_in_paragraph.append(f"pict:{embed_id}")
                            self.logger.debug(f"–ù–∞–π–¥–µ–Ω–æ pict –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {docx_idx}")
                
                # 4. –ò—â–µ–º –≤ run —ç–ª–µ–º–µ–Ω—Ç–∞—Ö (–¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –ø—Ä–æ–≤–µ—Ä–∫–∞)
                runs = paragraph.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                for run in runs:
                    # –ò—â–µ–º –≤—Å–µ —ç–ª–µ–º–µ–Ω—Ç—ã —Å embed –∞—Ç—Ä–∏–±—É—Ç–æ–º
                    embeds = run.findall('.//*[@r:embed]', namespaces={'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
                    for embed in embeds:
                        embed_id = embed.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id not in image_positions:
                            image_positions[embed_id] = docx_idx
                            images_in_paragraph.append(f"run:{embed_id}")
                            self.logger.debug(f"–ù–∞–π–¥–µ–Ω–æ run –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {docx_idx}")
                
                # 5. –î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–π –ø–æ–∏—Å–∫ –ø–æ –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–º –∞—Ç—Ä–∏–±—É—Ç–∞–º
                # –ò—â–µ–º —ç–ª–µ–º–µ–Ω—Ç—ã —Å r:id (alternative relationship format)
                alt_images = paragraph.findall('.//*[@r:id]', namespaces={'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'})
                for alt_img in alt_images:
                    embed_id = alt_img.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if embed_id and embed_id not in image_positions:
                        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ —ç—Ç–æ –¥–µ–π—Å—Ç–≤–∏—Ç–µ–ª—å–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
                        if self._is_image_relationship(embed_id):
                            image_positions[embed_id] = docx_idx
                            images_in_paragraph.append(f"alt:{embed_id}")
                            self.logger.debug(f"–ù–∞–π–¥–µ–Ω–æ –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω–æ–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {docx_idx}")
                
                # –õ–æ–≥–∏—Ä—É–µ–º –Ω–∞–π–¥–µ–Ω–Ω—ã–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–µ
                if images_in_paragraph:
                    self.logger.info(f"üìç –ü–∞—Ä–∞–≥—Ä–∞—Ñ XML[{xml_idx}] -> DOCX[{docx_idx}]: –Ω–∞–π–¥–µ–Ω–æ {len(images_in_paragraph)} –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {', '.join(images_in_paragraph)}")
                    print(f"üìç –ü–∞—Ä–∞–≥—Ä–∞—Ñ XML[{xml_idx}] -> DOCX[{docx_idx}]: {len(images_in_paragraph)} –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π - {', '.join(images_in_paragraph)}")
                    
            # –§–∏–Ω–∞–ª—å–Ω–∞—è —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞
            self.logger.info(f"üéØ –ò–¢–û–ì–û –Ω–∞–π–¥–µ–Ω–æ –ø–æ–∑–∏—Ü–∏–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {len(image_positions)}")
            print(f"üéØ –ò–¢–û–ì–û –Ω–∞–π–¥–µ–Ω–æ –ø–æ–∑–∏—Ü–∏–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {len(image_positions)}")
            
            # –î–µ—Ç–∞–ª—å–Ω—ã–π –≤—ã–≤–æ–¥ –≤—Å–µ—Ö –Ω–∞–π–¥–µ–Ω–Ω—ã—Ö –ø–æ–∑–∏—Ü–∏–π
            for embed_id, para_idx in sorted(image_positions.items(), key=lambda x: x[1]):
                self.logger.info(f"  üìå –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {para_idx}")
                print(f"  üìå –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {embed_id} -> –ø–∞—Ä–∞–≥—Ä–∞—Ñ {para_idx}")
                            
        except Exception as e:
            self.logger.error(f"‚ùå –û—à–∏–±–∫–∞ –£–õ–£–ß–®–ï–ù–ù–û–ì–û –ø–∞—Ä—Å–∏–Ω–≥–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –¥–ª—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {e}")
            print(f"‚ùå –û—à–∏–±–∫–∞ –£–õ–£–ß–®–ï–ù–ù–û–ì–û –ø–∞—Ä—Å–∏–Ω–≥–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –¥–ª—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {e}")
            
        return image_positions
    
    def _find_rel_id_for_media(self, media_file: str, relationships: Dict[str, str]) -> Optional[str]:
        """–ù–∞—Ö–æ–¥–∏—Ç relationship ID –¥–ª—è –º–µ–¥–∏–∞ —Ñ–∞–π–ª–∞"""
        media_filename = os.path.basename(media_file)
        
        # –ü–æ–ø—ã—Ç–∫–∞ —Ç–æ—á–Ω–æ–≥–æ —Å–æ–≤–ø–∞–¥–µ–Ω–∏—è
        for rel_id, target in relationships.items():
            if target.endswith(media_filename):
                return rel_id
        
        # –ü–æ–ø—ã—Ç–∫–∞ —Å–æ–≤–ø–∞–¥–µ–Ω–∏—è –±–µ–∑ "media/" –ø—Ä–µ—Ñ–∏–∫—Å–∞
        for rel_id, target in relationships.items():
            if target.endswith(media_filename) or target.endswith(media_file):
                return rel_id
        
        # –ü–æ–ø—ã—Ç–∫–∞ —Å–æ–≤–ø–∞–¥–µ–Ω–∏—è –ø–æ –∏–º–µ–Ω–∏ —Ñ–∞–π–ª–∞ –±–µ–∑ —Ä–∞—Å—à–∏—Ä–µ–Ω–∏—è
        media_name_without_ext = os.path.splitext(media_filename)[0]
        for rel_id, target in relationships.items():
            target_name = os.path.splitext(os.path.basename(target))[0]
            if target_name == media_name_without_ext:
                return rel_id
        
        self.logger.warning(f"–ù–µ –Ω–∞–π–¥–µ–Ω relationship ID –¥–ª—è –º–µ–¥–∏–∞ —Ñ–∞–π–ª–∞: {media_file}")
        return None
    
    def _get_image_dimensions(self, image_data: bytes) -> Tuple[Optional[float], Optional[float]]:
        """–ü–æ–ª—É—á–∞–µ—Ç —Ä–∞–∑–º–µ—Ä—ã –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –¥—é–π–º–∞—Ö"""
        try:
            with Image.open(io.BytesIO(image_data)) as img:
                # –ü–æ–ª—É—á–∞–µ–º —Ä–∞–∑–º–µ—Ä—ã –≤ –ø–∏–∫—Å–µ–ª—è—Ö
                width_px, height_px = img.size
                
                # –ü–æ–ª—É—á–∞–µ–º DPI (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é 96)
                dpi = img.info.get('dpi', (96, 96))
                if isinstance(dpi, tuple):
                    dpi_x, dpi_y = dpi
                else:
                    dpi_x = dpi_y = dpi
                
                # –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –≤ –¥—é–π–º—ã
                width_inches = width_px / dpi_x
                height_inches = height_px / dpi_y
                
                return width_inches, height_inches
                
        except Exception as e:
            self.logger.warning(f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ–ø—Ä–µ–¥–µ–ª–∏—Ç—å —Ä–∞–∑–º–µ—Ä—ã –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {e}")
            return None, None
    
    def _detect_image_format(self, image_data: bytes) -> str:
        """–û–ø—Ä–µ–¥–µ–ª—è–µ—Ç —Ñ–æ—Ä–º–∞—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –ø–æ binary data"""
        try:
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å–∏–≥–Ω–∞—Ç—É—Ä—ã —Ñ–∞–π–ª–æ–≤
            if image_data.startswith(b'\x89PNG'):
                return 'png'
            elif image_data.startswith(b'\xFF\xD8\xFF'):
                return 'jpeg'
            elif image_data.startswith(b'GIF87a') or image_data.startswith(b'GIF89a'):
                return 'gif'
            elif image_data.startswith(b'BM'):
                return 'bmp'
            elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[:12]:
                return 'webp'
            else:
                # –ü–æ–ø—ã—Ç–∫–∞ –æ–ø—Ä–µ–¥–µ–ª–∏—Ç—å —á–µ—Ä–µ–∑ PIL
                with Image.open(io.BytesIO(image_data)) as img:
                    return img.format.lower()
                    
        except Exception as e:
            self.logger.warning(f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ–ø—Ä–µ–¥–µ–ª–∏—Ç—å —Ñ–æ—Ä–º–∞—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {e}")
            return 'unknown'
    
    def insert_images_into_document(self, document: Document, original_docx_path: str) -> bool:
        """
        –í—Å—Ç–∞–≤–ª—è–µ—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –Ω–æ–≤—ã–π –¥–æ–∫—É–º–µ–Ω—Ç
        
        Args:
            document: –¶–µ–ª–µ–≤–æ–π –¥–æ–∫—É–º–µ–Ω—Ç
            original_docx_path: –ü—É—Ç—å –∫ –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–º—É –¥–æ–∫—É–º–µ–Ω—Ç—É
            
        Returns:
            True –µ—Å–ª–∏ —É—Å–ø–µ—à–Ω–æ, False –∏–Ω–∞—á–µ
        """
        try:
            # –°–Ω–∞—á–∞–ª–∞ –∏–∑–≤–ª–µ–∫–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –µ—Å–ª–∏ –µ—â—ë –Ω–µ —Å–¥–µ–ª–∞–ª–∏
            if not self.images:
                self.extract_images_from_docx(original_docx_path)
            
            # –í—Å—Ç–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–∏–µ –ø–æ–∑–∏—Ü–∏–∏
            for image_info in self.images:
                success = self._insert_single_image(document, image_info)
                if not success:
                    self.logger.warning(f"–ù–µ —É–¥–∞–ª–æ—Å—å –≤—Å—Ç–∞–≤–∏—Ç—å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {image_info.image_id}")
            
            return True
            
        except Exception as e:
            self.logger.error(f"–û—à–∏–±–∫–∞ –≤—Å—Ç–∞–≤–∫–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {e}")
            return False
    
    def _insert_single_image(self, document: Document, image_info: ImageInfo) -> bool:
        """–í—Å—Ç–∞–≤–ª—è–µ—Ç –æ–¥–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –≤ –¥–æ–∫—É–º–µ–Ω—Ç"""
        try:
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –ø–æ–∑–∏—Ü–∏—é –¥–ª—è –≤—Å—Ç–∞–≤–∫–∏
            target_paragraph_index = image_info.paragraph_index
            
            if target_paragraph_index is None or target_paragraph_index >= len(document.paragraphs):
                # –í—Å—Ç–∞–≤–ª—è–µ–º –≤ –∫–æ–Ω–µ—Ü –¥–æ–∫—É–º–µ–Ω—Ç–∞
                paragraph = document.add_paragraph()
            else:
                # –ü–æ–ª—É—á–∞–µ–º —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ –∏–ª–∏ —Å–æ–∑–¥–∞–µ–º –Ω–æ–≤—ã–π —Ä—è–¥–æ–º
                if target_paragraph_index < len(document.paragraphs):
                    # –í—Å—Ç–∞–≤–ª—è–µ–º –ø–æ—Å–ª–µ —É–∫–∞–∑–∞–Ω–Ω–æ–≥–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞
                    target_para = document.paragraphs[target_paragraph_index]
                    paragraph = document.add_paragraph()
                else:
                    paragraph = document.add_paragraph()
            
            # –ü–æ–ª—É—á–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –∏–∑ –≤—Ä–µ–º–µ–Ω–Ω–æ–π –ø–∞–ø–∫–∏
            temp_path = os.path.join(self.temp_dir, f"{image_info.image_id}.{image_info.image_format}")
            
            if not os.path.exists(temp_path):
                self.logger.warning(f"–§–∞–π–ª –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –Ω–µ –Ω–∞–π–¥–µ–Ω: {temp_path}")
                return False
            
            # –í—Å—Ç–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
            run = paragraph.add_run()
            
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ä–∞–∑–º–µ—Ä—ã
            if image_info.width and image_info.height:
                # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º —Ä–∞–∑–º–µ—Ä—ã
                max_width = 6.0  # –º–∞–∫—Å–∏–º—É–º 6 –¥—é–π–º–æ–≤
                max_height = 8.0  # –º–∞–∫—Å–∏–º—É–º 8 –¥—é–π–º–æ–≤
                
                width = min(image_info.width, max_width)
                height = min(image_info.height, max_height)
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø—Ä–æ–ø–æ—Ä—Ü–∏–∏
                if width / image_info.width < height / image_info.height:
                    height = width * (image_info.height / image_info.width)
                else:
                    width = height * (image_info.width / image_info.height)
                    
                run.add_picture(temp_path, width=Inches(width), height=Inches(height))
            else:
                # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Ä–∞–∑–º–µ—Ä –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
                run.add_picture(temp_path, width=Inches(4.0))
            
            self.logger.info(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ {image_info.image_id} —É—Å–ø–µ—à–Ω–æ –≤—Å—Ç–∞–≤–ª–µ–Ω–æ")
            return True
            
        except Exception as e:
            self.logger.warning(f"–û—à–∏–±–∫–∞ –≤—Å—Ç–∞–≤–∫–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è {image_info.image_id}: {e}")
            return False
    
    def cleanup_temp_files(self):
        """–û—á–∏—â–∞–µ—Ç –≤—Ä–µ–º–µ–Ω–Ω—ã–µ —Ñ–∞–π–ª—ã"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                import shutil
                shutil.rmtree(self.temp_dir)
                self.temp_dir = None
                self.logger.info("–í—Ä–µ–º–µ–Ω–Ω—ã–µ —Ñ–∞–π–ª—ã –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –æ—á–∏—â–µ–Ω—ã")
            except Exception as e:
                self.logger.error(f"–û—à–∏–±–∫–∞ –æ—á–∏—Å—Ç–∫–∏ –≤—Ä–µ–º–µ–Ω–Ω—ã—Ö —Ñ–∞–π–ª–æ–≤: {e}")
    
    def get_image_statistics(self) -> Dict[str, Any]:
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π"""
        if not self.images:
            return {'total_images': 0}
        
        formats = {}
        total_size = 0
        
        for image in self.images:
            # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º —Ñ–æ—Ä–º–∞—Ç—ã
            formats[image.image_format] = formats.get(image.image_format, 0) + 1
            
            # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º –æ–±—â–∏–π —Ä–∞–∑–º–µ—Ä
            total_size += len(image.image_data)
        
        return {
            'total_images': len(self.images),
            'formats': formats,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'average_size_kb': round((total_size / len(self.images)) / 1024, 2) if self.images else 0
        }
    
    def _is_image_relationship(self, rel_id: str) -> bool:
        """
        –ü—Ä–æ–≤–µ—Ä—è–µ—Ç, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ relationship ID –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ–º
        
        Args:
            rel_id: ID relationship –¥–ª—è –ø—Ä–æ–≤–µ—Ä–∫–∏
            
        Returns:
            True –µ—Å–ª–∏ —ç—Ç–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ, False –∏–Ω–∞—á–µ
        """
        # –†–∞—Å—à–∏—Ä–µ–Ω–Ω–∞—è –ø—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞ –æ—Å–Ω–æ–≤–µ –∞–Ω–∞–ª–∏–∑–∞ relationship
        # –ü–æ–∫–∞ –≤–æ–∑–≤—Ä–∞—â–∞–µ–º True –¥–ª—è –≤—Å–µ—Ö, –Ω–æ –º–æ–∂–Ω–æ –¥–æ–±–∞–≤–∏—Ç—å —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏—é –ø–æ —Ç–∏–ø—É
        self.logger.debug(f"üîç –ü—Ä–æ–≤–µ—Ä–∫–∞ relationship {rel_id} –Ω–∞ –ø—Ä–µ–¥–º–µ—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è")
        return True
    
    def get_detailed_extraction_log(self) -> str:
        """
        –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –¥–µ—Ç–∞–ª—å–Ω—ã–π –ª–æ–≥ –ø—Ä–æ—Ü–µ—Å—Å–∞ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
        
        Returns:
            –°—Ç—Ä–æ–∫–∞ —Å –¥–µ—Ç–∞–ª—å–Ω—ã–º –ª–æ–≥–æ–º
        """
        if not self.images:
            return "–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –Ω–µ –∏–∑–≤–ª–µ—á–µ–Ω—ã"
            
        log_lines = [
            f"üìä –î–ï–¢–ê–õ–¨–ù–´–ô –û–¢–ß–ï–¢ –ü–û –ò–ó–í–õ–ï–ß–ï–ù–ò–Æ –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô",
            f"=" * 50,
            f"–í—Å–µ–≥–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {len(self.images)}",
            f"–í—Ä–µ–º–µ–Ω–Ω–∞—è –ø–∞–ø–∫–∞: {self.temp_dir}",
            f"",
            f"üìã –°–ü–ò–°–û–ö –ò–ó–û–ë–†–ê–ñ–ï–ù–ò–ô:"
        ]
        
        positioned_count = 0
        unpositioned_count = 0
        
        for i, img in enumerate(self.images, 1):
            status = "‚úÖ –ü–û–ó–ò–¶–ò–û–ù–ò–†–û–í–ê–ù–û" if img.paragraph_index is not None else "‚ùì –ë–ï–ó –ü–û–ó–ò–¶–ò–ò"
            if img.paragraph_index is not None:
                positioned_count += 1
            else:
                unpositioned_count += 1
                
            log_lines.append(f"  {i}. {img.image_id}")
            log_lines.append(f"     –§–∞–π–ª: {img.filename}")
            log_lines.append(f"     –§–æ—Ä–º–∞—Ç: {img.image_format}")
            log_lines.append(f"     –†–∞–∑–º–µ—Ä: {img.width}x{img.height} –¥—é–π–º–æ–≤" if img.width and img.height else "     –†–∞–∑–º–µ—Ä: –Ω–µ–∏–∑–≤–µ—Å—Ç–µ–Ω")
            log_lines.append(f"     –ü–æ–∑–∏—Ü–∏—è: {img.paragraph_index}" if img.paragraph_index is not None else "     –ü–æ–∑–∏—Ü–∏—è: –Ω–µ –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∞")
            log_lines.append(f"     Relationship: {img.rel_id}")
            log_lines.append(f"     –°—Ç–∞—Ç—É—Å: {status}")
            log_lines.append("")
        
        log_lines.extend([
            f"üìà –°–¢–ê–¢–ò–°–¢–ò–ö–ê:",
            f"  –ü–æ–∑–∏—Ü–∏–æ–Ω–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö: {positioned_count}",
            f"  –ë–µ–∑ –ø–æ–∑–∏—Ü–∏–π: {unpositioned_count}",
            f"  –£—Å–ø–µ—à–Ω–æ—Å—Ç—å: {(positioned_count/len(self.images)*100):.1f}%",
        ])
        
        return "\n".join(log_lines) 