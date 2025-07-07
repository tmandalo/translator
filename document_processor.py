"""
Модуль для обработки .docx документов
"""

import os
import re
import traceback
import xml.etree.ElementTree as ET
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import time

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table

from improved_image_processor import ImprovedImageProcessor, ImageElement, ImageInfo
from image_adapter import ImageAdapter
from formatting_processor import FormattingProcessor
from translator import DocumentTranslator, TranslationResult


class TranslationProgress:
    """Класс для отслеживания прогресса перевода"""
    
    def __init__(self, total_items):
        self.total_items = total_items
        self.start_time = time.time()
        
    def __enter__(self):
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        pass
        
    def update(self, current_item, total_items, show_details=False):
        """Обновляет и отображает прогресс"""
        elapsed = time.time() - self.start_time
        percentage = (current_item / total_items) * 100 if total_items > 0 else 0
        
        if show_details:
            print(f"    Прогресс: {current_item}/{total_items} ({percentage:.1f}%) - {elapsed:.1f}с")


@dataclass
class DocumentElement:
    """Класс для хранения элемента документа"""
    element_type: str  # 'paragraph', 'table', 'header', 'footer', 'image'
    content: str
    original_element: Any
    index: int
    style: Optional[str] = None
    formatting: Optional[Dict[str, Any]] = None
    image_element: Optional[ImageElement] = None  # Для хранения информации об изображении


class DocumentProcessor:
    """Класс для обработки .docx документов"""
    
    def __init__(self):
        self.document = None
        self.elements: List[DocumentElement] = []
        self.improved_image_processor = ImprovedImageProcessor()
        self.images: List[ImageElement] = []
        self.file_path = None
        self.formatting_processor = FormattingProcessor()
        self.translator = DocumentTranslator()
        
        # СИСТЕМА ОТСЛЕЖИВАНИЯ ПОЗИЦИЙ
        self.position_tracker = {
            'extraction_stage': {},      # Позиции после извлечения из XML
            'validation_stage': {},      # Позиции после валидации
            'positioning_stage': {},     # Позиции после интеллектуального позиционирования
            'insertion_stage': {},       # Позиции в итоговом документе
            'position_history': [],      # История изменений позиций
            'tracking_enabled': True
        }
    
    def load_document(self, file_path: str) -> bool:
        """
        Загружает документ из файла
        
        Args:
            file_path: Путь к .docx файлу
            
        Returns:
            True если загрузка успешна, False иначе
        """
        try:
            self.document = Document(file_path)
            self.file_path = file_path
            self.elements = []
            return True
        except Exception as e:
            print(f"Ошибка загрузки документа: {e}")
            return False
    
    def process_and_translate(self) -> Optional[Document]:
        """
        ФИНАЛЬНАЯ ВЕРСИЯ: Главный метод, который выполняет поэлементную реконструкцию
        документа с переводом, сохраняя всю структуру.
        """
        if not self.document:
            print("❌ Документ не загружен.")
            return None

        # 1. Извлекаем информацию об изображениях и их позициях
        print("🔍 Шаг 1: Извлечение информации об изображениях...")
        image_infos = self.improved_image_processor.extract_images_from_docx(self.file_path)
        self.images = ImageAdapter.convert_list_to_image_elements(image_infos)
        
        images_by_paragraph = {}
        for img in self.images:
            if img.paragraph_index is not None:
                if img.paragraph_index not in images_by_paragraph:
                    images_by_paragraph[img.paragraph_index] = []
                images_by_paragraph[img.paragraph_index].append(img)
        
        print(f"✅ Найдено {len(self.images)} изображений, распределено по {len(images_by_paragraph)} параграфам.")

        # 2. Создаем новый, пустой документ для результата
        new_doc = Document()
        
        # 3. Итерируемся по КАЖДОМУ параграфу оригинального документа
        print("\n🔍 Шаг 2: Поэлементная реконструкция и перевод документа...")
        total_paragraphs = len(self.document.paragraphs)
        
        with TranslationProgress(total_paragraphs) as progress:
            for i, p in enumerate(self.document.paragraphs):
                
                # A. Вставляем изображения, которые идут ПЕРЕД этим параграфом
                if i in images_by_paragraph:
                    for image_element in sorted(images_by_paragraph[i], key=lambda img: img.image_id):
                        self._insert_image_with_smart_positioning(new_doc, image_element, i)
                        print(f"🖼️  Изображение {image_element.image_id} вставлено перед параграфом {i}")

                # B. Обрабатываем сам параграф
                if p.text.strip():
                    # Если есть текст - переводим
                    print(f"  Переводим параграф {i+1}/{total_paragraphs}...")
                    result = self.translator.api_translator.translate_text(p.text)
                    if result.success:
                        para_formatting = self._extract_paragraph_formatting(p)
                        new_para = new_doc.add_paragraph()
                        self._apply_advanced_formatting(new_para, p.text, result.translated_text, para_formatting)
                    else:
                        new_doc.add_paragraph(f"[ОШИБКА ПЕРЕВОДА] {p.text}")
                else:
                    # Если параграф пустой - просто добавляем пустой параграф для сохранения верстки
                    new_doc.add_paragraph()
                
                progress.update(i + 1, total_paragraphs, True)
        
        # TODO: Добавить такую же поэлементную обработку для таблиц, если требуется.

        print("\n✅ Реконструкция документа завершена.")
        return new_doc
    
    def _validate_and_correct_image_positions(self, images: List[ImageElement]) -> List[ImageElement]:
        """
        УЛУЧШЕННАЯ валидация и коррекция позиций изображений
        
        Args:
            images: Список изображений для валидации
            
        Returns:
            Список изображений с исправленными позициями
        """
        if not images or not self.document:
            return images
            
        total_paragraphs = len(self.document.paragraphs)
        print(f"🔍 ВАЛИДАЦИЯ: Проверяем {len(images)} изображений против {total_paragraphs} параграфов")
        
        # Создаем карту значимых параграфов для лучшего сопоставления
        significant_paragraphs = []
        for i, para in enumerate(self.document.paragraphs):
            has_text = para.text.strip()
            has_images = any(run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') for run in para.runs)
            
            if has_text or has_images:
                significant_paragraphs.append({
                    'index': i,
                    'text_preview': para.text[:50] + '...' if len(para.text) > 50 else para.text,
                    'has_text': has_text,
                    'has_images': has_images
                })
        
        print(f"🔍 ВАЛИДАЦИЯ: Найдено {len(significant_paragraphs)} значимых параграфов")
        
        # Статистика валидации
        stats = {
            'valid_positions': 0,
            'invalid_positions': 0,
            'corrected_positions': 0,
            'distributed_positions': 0,
            'end_positions': 0
        }
        
        corrected_images = []
        
        for img_idx, image in enumerate(images):
            original_position = image.paragraph_index
            
            # === ЭТАП 1: БАЗОВАЯ ВАЛИДАЦИЯ ===
            if image.paragraph_index is None:
                print(f"❓ ВАЛИДАЦИЯ: Изображение {image.image_id} не имеет позиции")
                stats['invalid_positions'] += 1
            elif image.paragraph_index < 0:
                print(f"⚠️  ВАЛИДАЦИЯ: Изображение {image.image_id} имеет отрицательную позицию {image.paragraph_index}")
                image.paragraph_index = None
                stats['invalid_positions'] += 1
            elif image.paragraph_index >= total_paragraphs:
                print(f"⚠️  ВАЛИДАЦИЯ: Изображение {image.image_id} имеет позицию {image.paragraph_index} превышающую количество параграфов ({total_paragraphs})")
                image.paragraph_index = None
                stats['invalid_positions'] += 1
            else:
                # Проверяем, является ли позиция значимой
                target_para = self.document.paragraphs[image.paragraph_index]
                has_meaningful_content = (target_para.text.strip() or 
                                        any(run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') for run in target_para.runs))
                
                if has_meaningful_content:
                    print(f"✅ ВАЛИДАЦИЯ: Изображение {image.image_id} имеет валидную позицию {image.paragraph_index}")
                    stats['valid_positions'] += 1
                else:
                    print(f"⚠️  ВАЛИДАЦИЯ: Изображение {image.image_id} привязано к пустому параграфу {image.paragraph_index}")
                    # Попытка найти ближайший значимый параграф
                    corrected_position = self._find_nearest_significant_paragraph(image.paragraph_index, significant_paragraphs)
                    if corrected_position is not None:
                        print(f"🔧 КОРРЕКЦИЯ: Изображение {image.image_id} перемещено с позиции {image.paragraph_index} на {corrected_position}")
                        image.paragraph_index = corrected_position
                        stats['corrected_positions'] += 1
                    else:
                        print(f"❌ КОРРЕКЦИЯ: Не удалось найти подходящую позицию для изображения {image.image_id}")
                        image.paragraph_index = None
                        stats['invalid_positions'] += 1
            
            # === ЭТАП 2: ИНТЕЛЛЕКТУАЛЬНАЯ КОРРЕКЦИЯ ===
            if image.paragraph_index is None and original_position is not None:
                # Попытка исправить позицию на основе анализа
                corrected_position = self._intelligent_position_correction(original_position, total_paragraphs, significant_paragraphs)
                if corrected_position is not None:
                    print(f"🧠 УМНАЯ КОРРЕКЦИЯ: Изображение {image.image_id} получило позицию {corrected_position} (было {original_position})")
                    image.paragraph_index = corrected_position
                    stats['corrected_positions'] += 1
            
            corrected_images.append(image)
        
        # === ЭТАП 3: РАСПРЕДЕЛЕНИЕ ИЗОБРАЖЕНИЙ БЕЗ ПОЗИЦИЙ ===
        images_without_position = [img for img in corrected_images if img.paragraph_index is None]
        
        if images_without_position:
            print(f"🎯 РАСПРЕДЕЛЕНИЕ: Обрабатываем {len(images_without_position)} изображений без позиций")
            
            # Стратегия распределения
            distribution_strategy = self._determine_distribution_strategy(len(images_without_position), len(significant_paragraphs))
            print(f"🎯 РАСПРЕДЕЛЕНИЕ: Используем стратегию '{distribution_strategy}'")
            
            if distribution_strategy == 'distribute':
                # Распределяем изображения по документу
                distributed_count = self._distribute_images_intelligently(images_without_position, significant_paragraphs)
                stats['distributed_positions'] += distributed_count
                stats['end_positions'] += len(images_without_position) - distributed_count
            elif distribution_strategy == 'cluster':
                # Группируем изображения в определенных местах
                clustered_count = self._cluster_images_strategically(images_without_position, significant_paragraphs)
                stats['distributed_positions'] += clustered_count
                stats['end_positions'] += len(images_without_position) - clustered_count
            else:
                # Все изображения в конец (текущее поведение)
                stats['end_positions'] += len(images_without_position)
        
        # === ФИНАЛЬНАЯ СТАТИСТИКА ===
        print(f"\n📊 ИТОГИ ВАЛИДАЦИИ:")
        print(f"  ✅ Валидные позиции: {stats['valid_positions']}")
        print(f"  ❌ Невалидные позиции: {stats['invalid_positions']}")
        print(f"  🔧 Исправленные позиции: {stats['corrected_positions']}")
        print(f"  🎯 Распределенные позиции: {stats['distributed_positions']}")
        print(f"  📌 Позиции в конце: {stats['end_positions']}")
        
        return corrected_images
    
    def _perform_hybrid_validation(self):
        """
        ГИБРИДНАЯ ВАЛИДАЦИЯ: Дополнительная диагностика позиций изображений
        Проверяет соответствие между XML-парсингом и python-docx API
        """
        if not self.images or not self.document:
            return
            
        total_paragraphs = len(self.document.paragraphs)
        
        print(f"\n🔍 ГИБРИДНАЯ ВАЛИДАЦИЯ: Проверяем корректность позиций")
        print(f"📊 Python-docx видит: {total_paragraphs} параграфов")
        
        # Проверяем количество параграфов с XML стороны
        xml_paragraphs_count = None
        if hasattr(self.improved_image_processor, '_last_positions'):
            # Попытка получить количество XML параграфов из логов процессора
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(self.file_path, 'r') as docx_zip:
                    doc_content = docx_zip.read('word/document.xml')
                    root = ET.fromstring(doc_content)
                    body = root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
                    
                    if body is not None:
                        xml_paragraphs = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                        xml_paragraphs_count = len(xml_paragraphs)
                        print(f"📊 XML-парсер видит: {xml_paragraphs_count} параграфов")
                        
            except Exception as e:
                print(f"⚠️  Не удалось получить количество XML параграфов: {e}")
        
        # Валидация 1: Проверка соответствия количества параграфов
        validation_issues = []
        
        if xml_paragraphs_count is not None:
            if xml_paragraphs_count != total_paragraphs:
                issue = f"Несоответствие количества параграфов: XML={xml_paragraphs_count}, python-docx={total_paragraphs}"
                validation_issues.append(issue)
                print(f"⚠️  {issue}")
            else:
                print(f"✅ Количество параграфов соответствует: {total_paragraphs}")
        
        # Валидация 2: Проверка позиций изображений
        valid_positions = 0
        invalid_positions = 0
        out_of_range_positions = 0
        
        for image in self.images:
            if image.paragraph_index is None:
                invalid_positions += 1
                print(f"⚠️  Изображение {image.image_id}: позиция не определена (None)")
            elif image.paragraph_index < 0:
                invalid_positions += 1
                validation_issues.append(f"Изображение {image.image_id} имеет отрицательную позицию: {image.paragraph_index}")
                print(f"⚠️  Изображение {image.image_id}: отрицательная позиция {image.paragraph_index}")
            elif image.paragraph_index >= total_paragraphs:
                out_of_range_positions += 1
                validation_issues.append(f"Изображение {image.image_id} имеет позицию {image.paragraph_index}, превышающую количество параграфов ({total_paragraphs})")
                print(f"❌ Изображение {image.image_id}: позиция {image.paragraph_index} превышает максимум ({total_paragraphs-1})")
            else:
                valid_positions += 1
                print(f"✅ Изображение {image.image_id}: валидная позиция {image.paragraph_index}")
        
        # Валидация 3: Статистика и рекомендации
        total_images = len(self.images)
        success_rate = (valid_positions / total_images * 100) if total_images > 0 else 0
        
        print(f"\n📊 РЕЗУЛЬТАТЫ ГИБРИДНОЙ ВАЛИДАЦИИ:")
        print(f"  ✅ Валидные позиции: {valid_positions}/{total_images} ({success_rate:.1f}%)")
        print(f"  ❌ Невалидные позиции: {invalid_positions}")
        print(f"  🚫 Позиции вне диапазона: {out_of_range_positions}")
        
        # Предупреждения и рекомендации
        if validation_issues:
            print(f"\n⚠️  ОБНАРУЖЕНЫ ПРОБЛЕМЫ ({len(validation_issues)}):")
            for i, issue in enumerate(validation_issues, 1):
                print(f"   {i}. {issue}")
                
            if success_rate < 50:
                print(f"\n🚨 КРИТИЧЕСКОЕ ПРЕДУПРЕЖДЕНИЕ: Успешность позиционирования составляет всего {success_rate:.1f}%")
                print(f"   Рекомендуется проверить алгоритм извлечения позиций в improved_image_processor.py")
            elif success_rate < 80:
                print(f"\n⚠️  ПРЕДУПРЕЖДЕНИЕ: Успешность позиционирования {success_rate:.1f}% ниже ожидаемой")
                print(f"   Возможны проблемы с форматом документа или edge cases")
        else:
            print(f"✅ Все проверки пройдены успешно!")
        
        # Валидация 4: Проверка memory в update_memory tool при обнаружении критических проблем
        if success_rate < 30 and len(validation_issues) > 3:
            print(f"\n🔧 СИСТЕМА САМОДИАГНОСТИКИ: Обнаружены серьезные проблемы позиционирования")
            print(f"   Возможно, требуется дополнительная отладка алгоритма извлечения изображений")
        
        print(f"─" * 60)
    
    def _find_nearest_significant_paragraph(self, target_index: int, significant_paragraphs: List[dict]) -> Optional[int]:
        """Находит ближайший значимый параграф к заданному индексу"""
        if not significant_paragraphs:
            return None
            
        # Ищем ближайший параграф по индексу
        best_distance = float('inf')
        best_index = None
        
        for para_info in significant_paragraphs:
            distance = abs(para_info['index'] - target_index)
            if distance < best_distance:
                best_distance = distance
                best_index = para_info['index']
        
        return best_index
    
    def _intelligent_position_correction(self, original_position: int, total_paragraphs: int, significant_paragraphs: List[dict]) -> Optional[int]:
        """Интеллектуальная коррекция позиции изображения"""
        if not significant_paragraphs:
            return None
            
        # Стратегия 1: Если позиция слишком большая, пропорционально уменьшаем
        if original_position >= total_paragraphs:
            proportion = original_position / total_paragraphs
            if proportion <= 2.0:  # Не более чем в 2 раза больше
                # Масштабируем к количеству значимых параграфов
                corrected_index = int((original_position / total_paragraphs) * len(significant_paragraphs))
                if corrected_index < len(significant_paragraphs):
                    return significant_paragraphs[corrected_index]['index']
        
        # Стратегия 2: Если позиция близка к концу, используем один из последних параграфов
        if original_position >= total_paragraphs * 0.8:
            last_third = significant_paragraphs[-len(significant_paragraphs)//3:] if len(significant_paragraphs) > 3 else significant_paragraphs
            if last_third:
                return last_third[0]['index']
        
        # Стратегия 3: Если позиция в начале, используем один из первых параграфов
        if original_position <= total_paragraphs * 0.2:
            first_third = significant_paragraphs[:len(significant_paragraphs)//3] if len(significant_paragraphs) > 3 else significant_paragraphs
            if first_third:
                return first_third[-1]['index']
        
        return None
    
    def _determine_distribution_strategy(self, images_count: int, paragraphs_count: int) -> str:
        """Определяет стратегию распределения изображений без позиций"""
        if images_count <= 2:
            return 'end'  # Мало изображений - в конец
        elif images_count <= paragraphs_count // 3:
            return 'distribute'  # Средне изображений - распределяем
        elif images_count <= paragraphs_count // 2:
            return 'cluster'  # Много изображений - группируем
        else:
            return 'end'  # Очень много изображений - в конец
    
    def _distribute_images_intelligently(self, images: List[ImageElement], significant_paragraphs: List[dict]) -> int:
        """Интеллектуально распределяет изображения по документу"""
        if not images or not significant_paragraphs:
            return 0
            
        distributed_count = 0
        
        # Вычисляем позиции для распределения
        step = len(significant_paragraphs) // (len(images) + 1)
        if step < 1:
            step = 1
            
        for i, image in enumerate(images):
            target_position = min((i + 1) * step, len(significant_paragraphs) - 1)
            if target_position < len(significant_paragraphs):
                image.paragraph_index = significant_paragraphs[target_position]['index']
                print(f"🎯 РАСПРЕДЕЛЕНИЕ: Изображение {image.image_id} размещено в позиции {image.paragraph_index}")
                distributed_count += 1
        
        return distributed_count
    
    def _cluster_images_strategically(self, images: List[ImageElement], significant_paragraphs: List[dict]) -> int:
        """Группирует изображения в стратегических местах документа"""
        if not images or not significant_paragraphs:
            return 0
            
        clustered_count = 0
        
        # Определяем точки кластеризации (начало, середина, конец)
        cluster_points = []
        if len(significant_paragraphs) > 10:
            cluster_points = [
                significant_paragraphs[len(significant_paragraphs)//4]['index'],  # Первая четверть
                significant_paragraphs[len(significant_paragraphs)//2]['index'],  # Середина
                significant_paragraphs[3*len(significant_paragraphs)//4]['index']  # Последняя четверть
            ]
        else:
            cluster_points = [
                significant_paragraphs[0]['index'],  # Начало
                significant_paragraphs[-1]['index']  # Конец
            ]
        
        # Распределяем изображения по кластерам
        for i, image in enumerate(images):
            if i < len(cluster_points):
                image.paragraph_index = cluster_points[i]
                print(f"🎯 КЛАСТЕРИЗАЦИЯ: Изображение {image.image_id} размещено в кластере на позиции {image.paragraph_index}")
                clustered_count += 1
        
        return clustered_count

    def extract_text_elements(self) -> List[DocumentElement]:
        """
        Извлекает текстовые элементы из документа
        
        Returns:
            Список элементов документа
        """
        if not self.document:
            return []
        
        elements = []
        element_index = 0
        
        # Сначала извлекаем все изображения (ТОЛЬКО улучшенный метод)
        if self.file_path:
            print(f"🔍 Используем УЛУЧШЕННЫЙ процессор изображений для файла: {self.file_path}")
            print(f"🔍 DocumentProcessor видит {len(self.document.paragraphs)} параграфов в документе")
            
            # Логируем этап извлечения
            self._log_image_processing_stage('extraction', {
                'file_path': self.file_path,
                'total_paragraphs': len(self.document.paragraphs)
            })
            
            image_infos = self.improved_image_processor.extract_images_from_docx(self.file_path)
            self.images = ImageAdapter.convert_list_to_image_elements(image_infos)
            print(f"🔍 Результат улучшенного процессора: {len(self.images)} изображений")
            
            # === ГИБРИДНАЯ ВАЛИДАЦИЯ (дополнительная защита) ===
            self._perform_hybrid_validation()
            
            # Трекинг позиций после извлечения
            self._track_image_positions('extraction', self.images, {'source': 'improved_processor'})
            
            # Логируем результаты извлечения
            self._log_image_processing_stage('extraction', {
                'file_path': self.file_path,
                'total_paragraphs': len(self.document.paragraphs),
                'images_found': len(self.images),
                'relationships_count': len(getattr(self.improved_image_processor, '_last_relationships', {})),
                'xml_positions_count': len(getattr(self.improved_image_processor, '_last_positions', {}))
            })
            
            # Валидируем и корректируем индексы изображений с помощью УЛУЧШЕННОЙ системы
            print(f"🔍 Используем УЛУЧШЕННУЮ валидацию позиций изображений")
            
            # Считаем статистику до валидации
            pre_validation_stats = {
                'total_images': len(self.images),
                'with_positions': len([img for img in self.images if img.paragraph_index is not None]),
                'without_positions': len([img for img in self.images if img.paragraph_index is None])
            }
            
            self.images = self._validate_and_correct_image_positions(self.images)
            
            # Трекинг позиций после валидации
            self._track_image_positions('validation', self.images, pre_validation_stats)
            
            # Считаем статистику после валидации
            post_validation_stats = {
                'total_images': len(self.images),
                'with_positions': len([img for img in self.images if img.paragraph_index is not None]),
                'without_positions': len([img for img in self.images if img.paragraph_index is None])
            }
            
            # Выводим детальный отчет по извлечению изображений
            if self.improved_image_processor and hasattr(self.improved_image_processor, 'get_detailed_extraction_log'):
                extraction_log = self.improved_image_processor.get_detailed_extraction_log()
                print(f"\n{extraction_log}\n")
        else:
            print("❌ ОШИБКА: Нет пути к файлу, не можем использовать улучшенный процессор!")
            # Без пути к файлу не можем извлечь изображения
            self.images = []
        
        # УЛУЧШЕННАЯ И НАДЕЖНАЯ система позиционирования изображений
        print(f"🖼️  Всего изображений: {len(self.images)}")
        
        # Создаем продвинутый маппинг изображений с анализом конфликтов
        images_by_paragraph = {}
        images_without_position = []
        positioning_conflicts = []
        
        # Анализируем все изображения и создаем карту позиций
        for image in self.images:
            if image.paragraph_index is not None and image.paragraph_index >= 0:
                # Проверяем конфликты позиций (несколько изображений на одной позиции)
                if image.paragraph_index not in images_by_paragraph:
                    images_by_paragraph[image.paragraph_index] = []
                images_by_paragraph[image.paragraph_index].append(image)
                
                # Логируем конфликты
                if len(images_by_paragraph[image.paragraph_index]) > 1:
                    positioning_conflicts.append(image.paragraph_index)
                
                print(f"🖼️  Изображение {image.image_id} привязано к параграфу {image.paragraph_index}")
            else:
                images_without_position.append(image)
                print(f"🖼️  Изображение {image.image_id} без позиции")
        
        # Логируем конфликты позиций
        if positioning_conflicts:
            print(f"⚠️  КОНФЛИКТЫ ПОЗИЦИЙ: {len(set(positioning_conflicts))} позиций с множественными изображениями")
            for pos in set(positioning_conflicts):
                print(f"   Позиция {pos}: {len(images_by_paragraph[pos])} изображений")
        
        # НОВАЯ СТРАТЕГИЯ: Интеллектуальное распределение элементов
        elements = []
        element_index = 0
        processed_images_count = 0
        
        # Подсчитываем статистику для стратегического распределения
        total_text_paragraphs = len([p for p in self.document.paragraphs if p.text.strip()])
        
        print(f"📊 СТРАТЕГИЯ РАСПРЕДЕЛЕНИЯ:")
        print(f"  • Параграфов с текстом: {total_text_paragraphs}")
        print(f"  • Изображений с позициями: {len(self.images) - len(images_without_position)}")
        print(f"  • Изображений без позиций: {len(images_without_position)}")
        
        # === ЭТАП 1: ОБРАБОТКА ПАРАГРАФОВ С ПРИВЯЗАННЫМИ ИЗОБРАЖЕНИЯМИ ===
        for paragraph_index, paragraph in enumerate(self.document.paragraphs):
            # СНАЧАЛА добавляем все изображения для этого параграфа
            if paragraph_index in images_by_paragraph:
                images_for_paragraph = images_by_paragraph[paragraph_index]
                
                # Сортируем изображения по ID для стабильного порядка
                images_for_paragraph.sort(key=lambda img: img.image_id)
                
                for image in images_for_paragraph:
                    image_element = DocumentElement(
                        element_type='image',
                        content=f"[IMAGE: {image.image_id}]",
                        original_element=paragraph,
                        index=element_index,
                        image_element=image
                    )
                    elements.append(image_element)
                    element_index += 1
                    processed_images_count += 1
                    print(f"✅ Изображение {image.image_id} добавлено ПЕРЕД параграфом {paragraph_index}")
            
            # ПОТОМ добавляем сам параграф (если есть текст)
            has_text = paragraph.text.strip()
            if has_text:
                element = DocumentElement(
                    element_type='paragraph',
                    content=paragraph.text,
                    original_element=paragraph,
                    index=element_index,
                    style=paragraph.style.name if paragraph.style else None,
                    formatting=self._extract_paragraph_formatting(paragraph)
                )
                elements.append(element)
                element_index += 1
        
        print(f"📊 ОБРАБОТАНО: {processed_images_count} изображений с определенными позициями")
        
        # === ЭТАП 2: ИНТЕЛЛЕКТУАЛЬНАЯ ОБРАБОТКА ИЗОБРАЖЕНИЙ БЕЗ ПОЗИЦИЙ ===
        if images_without_position:
            print(f"🎯 ИНТЕЛЛЕКТУАЛЬНОЕ РАСПРЕДЕЛЕНИЕ: {len(images_without_position)} изображений без позиций")
            
            # Определяем стратегию на основе анализа документа
            distribution_strategy = self._determine_smart_distribution_strategy(
                images_without_position, 
                total_text_paragraphs, 
                len(elements)
            )
            
            print(f"🎯 ВЫБРАННАЯ СТРАТЕГИЯ: {distribution_strategy}")
            
            # Применяем выбранную стратегию
            if distribution_strategy == 'strategic_insertion':
                elements = self._insert_images_strategically(elements, images_without_position, element_index)
            elif distribution_strategy == 'proportional_distribution':
                elements = self._distribute_images_proportionally(elements, images_without_position, element_index)
            elif distribution_strategy == 'chapter_clustering':
                elements = self._cluster_images_by_chapters(elements, images_without_position, element_index)
            else:  # 'end_placement'
                elements = self._place_images_at_end(elements, images_without_position, element_index)
        
        # === ЭТАП 3: ОБРАБОТКА ТАБЛИЦ ===
        for table in self.document.tables:
            table_text = self._extract_table_text(table)
            if table_text.strip():
                element = DocumentElement(
                    element_type='table',
                    content=table_text,
                    original_element=table,
                    index=element_index,
                    formatting=self._extract_table_formatting(table)
                )
                elements.append(element)
                element_index += 1
        
        self.elements = elements
        
        # Логируем финальную статистику позиционирования
        positioned_images = len([elem for elem in elements if elem.element_type == 'image' and elem.image_element and elem.image_element.paragraph_index is not None])
        unpositioned_images = len([elem for elem in elements if elem.element_type == 'image' and (not elem.image_element or elem.image_element.paragraph_index is None)])
        text_paragraphs = len([elem for elem in elements if elem.element_type == 'paragraph'])
        tables_count = len([elem for elem in elements if elem.element_type == 'table'])
        
        self._log_image_processing_stage('positioning', {
            'total_elements': len(elements),
            'positioned_images': positioned_images,
            'unpositioned_images': unpositioned_images,
            'text_paragraphs': text_paragraphs,
            'tables_count': tables_count
        })
        
        return elements
    
    def _extract_paragraph_formatting(self, paragraph: Paragraph) -> Dict[str, Any]:
        """Извлекает форматирование параграфа"""
        formatting = {
            'alignment': paragraph.alignment,
            'runs': []
        }
        
        for run in paragraph.runs:
            run_formatting = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color.rgb else None
            }
            formatting['runs'].append(run_formatting)
        
        return formatting
    
    def _extract_table_text(self, table: Table) -> str:
        """Извлекает текст из таблицы"""
        text_parts = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts)
    
    def _extract_table_formatting(self, table: Table) -> Dict[str, Any]:
        """Извлекает форматирование таблицы"""
        return {
            'rows': len(table.rows),
            'cols': len(table.columns),
            'style': table.style.name if table.style else None
        }
    
    def update_element_content(self, element_index: int, new_content: str) -> bool:
        """
        Обновляет содержимое элемента
        
        Args:
            element_index: Индекс элемента
            new_content: Новое содержимое
            
        Returns:
            True если обновление успешно, False иначе
        """
        if element_index >= len(self.elements):
            return False
        
        element = self.elements[element_index]
        
        try:
            if element.element_type == 'paragraph':
                self._update_paragraph_content(element.original_element, new_content)
            elif element.element_type == 'table':
                self._update_table_content(element.original_element, new_content)
            elif element.element_type == 'image':
                # Изображения не обновляются, только их описание
                pass
            
            # Обновляем содержимое в нашем списке
            element.content = new_content
            return True
            
        except Exception as e:
            print(f"Ошибка обновления элемента {element_index}: {e}")
            return False
    
    def _update_paragraph_content(self, paragraph: Paragraph, new_content: str):
        """Обновляет содержимое параграфа с сохранением форматирования"""
        # Сохраняем форматирование первого run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            
            # Очищаем параграф
            paragraph.clear()
            
            # Создаем новый run с сохраненным форматированием
            new_run = paragraph.add_run(new_content)
            
            # Применяем форматирование
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            if first_run.font.name:
                new_run.font.name = first_run.font.name
            if first_run.font.size:
                new_run.font.size = first_run.font.size
            if first_run.font.color.rgb:
                new_run.font.color.rgb = first_run.font.color.rgb
        else:
            # Если нет runs, просто добавляем текст
            paragraph.add_run(new_content)
    
    def _update_table_content(self, table: Table, new_content: str):
        """Обновляет содержимое таблицы"""
        # Разбиваем новый контент на строки
        rows_content = new_content.split('\n')
        
        for row_idx, row in enumerate(table.rows):
            if row_idx < len(rows_content):
                # Разбиваем строку на ячейки
                cells_content = rows_content[row_idx].split(' | ')
                
                for cell_idx, cell in enumerate(row.cells):
                    if cell_idx < len(cells_content):
                        # Обновляем содержимое ячейки
                        cell.text = cells_content[cell_idx].strip()
    
    def save_document(self, output_path: str) -> bool:
        """
        Сохраняет текущий документ в файл
        
        Args:
            output_path: Путь для сохранения
            
        Returns:
            True если сохранение успешно, False иначе
        """
        if not self.document:
            return False
        
        try:
            # Создаем директорию если она не существует
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            self.document.save(output_path)
            return True
            
        except Exception as e:
            print(f"Ошибка сохранения документа: {e}")
            return False
    
    def save_document_with_images(self, document: Document, output_path: str) -> bool:
        """
        Сохраняет документ с изображениями в файл
        
        Args:
            document: Документ для сохранения
            output_path: Путь для сохранения
            
        Returns:
            True если сохранение успешно, False иначе
        """
        try:
            # Создаем директорию если она не существует
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            document.save(output_path)
            return True
            
        except Exception as e:
            print(f"Ошибка сохранения документа с изображениями: {e}")
            return False
    
    def create_translated_document(self, translation_results: List[Any]) -> Optional[Document]:
        """
        ИСПРАВЛЕНО v2: Создает документ, корректно обрабатывая пустые строки для точной верстки.
        """
        try:
            new_document = Document()
            EMPTY_PARA_MARKER = "[[EMPTY_PARAGRAPH_MARKER]]"
            
            full_translated_text = '\n\n'.join(
                res.translated_text for res in translation_results if res.success
            )
            
            translated_paragraphs = list(filter(None, re.split(r'\n\s*\n', full_translated_text)))
            translated_paragraph_iterator = iter(translated_paragraphs)

            print(f"🔄 Обрабатываем {len(self.elements)} элементов документа.")
            print(f"📄 Получено {len(translated_paragraphs)} переведенных текстовых блоков для вставки.")

            for element_idx, element in enumerate(self.elements):
                if element.element_type == 'image':
                    if element.image_element:
                        self._insert_image_with_smart_positioning(new_document, element.image_element, element_idx)
                
                elif element.element_type == 'paragraph':
                    if not element.content.strip():
                        new_document.add_paragraph()
                        print(f"📄 Пустой параграф (элемент {element_idx}) сохранен для верстки.")
                    else:
                        try:
                            translated_text = next(translated_paragraph_iterator)
                            if translated_text.strip() and translated_text != EMPTY_PARA_MARKER:
                                self._create_translated_paragraph_with_context(
                                    new_document, element, translated_text, element_idx
                                )
                            else:
                                new_document.add_paragraph()
                        except StopIteration:
                            print(f"⚠️  Предупреждение: закончился переведенный текст на элементе {element_idx}.")
                
                elif element.element_type == 'table':
                    try:
                        translated_text_for_table = next(translated_paragraph_iterator)
                        self._add_translated_table(new_document, translated_text_for_table, element.formatting)
                    except StopIteration:
                        print(f"⚠️  Предупреждение: закончился переведенный текст для таблицы на элементе {element_idx}.")
            
            remaining_paragraphs = list(translated_paragraph_iterator)
            if remaining_paragraphs:
                print(f"⚠️  Предупреждение: {len(remaining_paragraphs)} переведенных параграфов остались неиспользованными. Вставляем их в конец.")
                for rem_para in remaining_paragraphs:
                    new_document.add_paragraph(rem_para)

            print(f"✅ Документ создан: {len(new_document.paragraphs)} параграфов, {len(new_document.inline_shapes)} изображений.")
            return new_document
            
        except Exception as e:
            print(f"❌ КРИТИЧЕСКАЯ Ошибка создания переведенного документа: {e}")
            traceback.print_exc()
            return None
    
    def _apply_advanced_formatting(self, paragraph: Paragraph, original_text: str, 
                                 translated_text: str, formatting_data: Dict[str, Any]):
        """
        Применяет КОНСЕРВАТИВНОЕ форматирование к параграфу для избежания проблем
        
        Args:
            paragraph: Параграф для форматирования
            original_text: Оригинальный текст
            translated_text: Переведенный текст
            formatting_data: Данные форматирования
        """
        try:
            # Если нет данных форматирования, просто добавляем текст
            if not formatting_data:
                paragraph.add_run(translated_text)
                return
            
            # Извлекаем сегменты форматирования из оригинального текста
            original_segments = self.formatting_processor.extract_formatting_segments(
                original_text, formatting_data
            )
            
            # Используем КОНСЕРВАТИВНОЕ сопоставление форматирования
            translated_segments = self.formatting_processor.map_conservative_formatting_to_translation(
                original_segments, original_text, translated_text
            )
            
            # Применяем форматирование к параграфу
            paragraph_alignment = formatting_data.get('alignment')
            success = self.formatting_processor.apply_formatting_to_paragraph(
                paragraph, translated_segments, paragraph_alignment
            )
            
            if not success:
                # Если не удалось применить форматирование, добавляем текст без форматирования
                paragraph.clear()
                paragraph.add_run(translated_text)
                
        except Exception as e:
            print(f"Ошибка применения консервативного форматирования: {e}")
            # В случае ошибки добавляем текст без форматирования
            paragraph.clear()
            paragraph.add_run(translated_text)
    
    def _insert_image_with_smart_positioning(self, document: Document, image_element: ImageElement, element_index: int) -> bool:
        """
        УЛУЧШЕННАЯ вставка изображения с умным позиционированием
        
        Args:
            document: Целевой документ
            image_element: Элемент изображения
            element_index: Индекс элемента в общем списке
            
        Returns:
            True если вставка успешна, False иначе
        """
        try:
            # Анализируем контекст для принятия решения о позиционировании
            positioning_context = self._analyze_image_positioning_context(element_index)
            
            # Получаем путь к временному файлу изображения
            temp_path = self._get_image_temp_path(image_element)
            if not temp_path:
                return False
            
            # Определяем стратегию вставки на основе контекста
            if positioning_context['use_existing_paragraph']:
                # Вставляем в существующий параграф
                target_paragraph = positioning_context['target_paragraph']
                success = self._insert_image_into_existing_paragraph(target_paragraph, image_element, temp_path)
            else:
                # Создаем новый параграф для изображения
                success = self._insert_image_into_new_paragraph(document, image_element, temp_path, positioning_context)
            
            if success:
                print(f"🖼️  УМНОЕ ПОЗИЦИОНИРОВАНИЕ: Изображение {image_element.image_id} вставлено по стратегии '{positioning_context['strategy']}'")
            
            return success
            
        except Exception as e:
            print(f"❌ Ошибка умной вставки изображения {image_element.image_id}: {e}")
            return False
    
    def _create_translated_paragraph_with_context(self, document: Document, element: DocumentElement, translated_text: str, element_index: int) -> Paragraph:
        """
        УЛУЧШЕННОЕ создание параграфа с учетом контекста и окружающих элементов
        
        Args:
            document: Целевой документ
            element: Исходный элемент
            translated_text: Переведенный текст
            element_index: Индекс элемента
            
        Returns:
            Созданный параграф
        """
        # Анализируем контекст параграфа
        context = self._analyze_paragraph_context(element_index)
        
        # Создаем параграф с учетом контекста
        if context['needs_spacing_before']:
            # Добавляем дополнительный отступ если нужно
            spacing_paragraph = document.add_paragraph()
            spacing_paragraph.add_run("")  # Пустой параграф для отступа
        
        # Создаем основной параграф
        paragraph = document.add_paragraph()
        
        # Применяем стиль параграфа
        if element.style:
            try:
                paragraph.style = element.style
            except:
                print(f"⚠️  Не удалось применить стиль '{element.style}' к параграфу")
        
        # Применяем детальное форматирование
        self._apply_advanced_formatting(
            paragraph, 
            element.content, 
            translated_text, 
            element.formatting
        )
        
        # Дополнительные настройки на основе контекста
        if context['is_title']:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif context['is_quote']:
            # Применяем отступы для цитат
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.right_indent = Inches(0.5)
        
        return paragraph
    
    def _analyze_image_positioning_context(self, element_index: int) -> Dict[str, Any]:
        """Анализирует контекст для определения стратегии позиционирования изображения"""
        context = {
            'strategy': 'new_paragraph',
            'use_existing_paragraph': False,
            'target_paragraph': None,
            'needs_spacing': True,
            'alignment': WD_ALIGN_PARAGRAPH.CENTER
        }
        
        # Анализируем соседние элементы
        prev_element = self.elements[element_index - 1] if element_index > 0 else None
        next_element = self.elements[element_index + 1] if element_index < len(self.elements) - 1 else None
        
        # Определяем стратегию на основе контекста
        if prev_element and prev_element.element_type == 'image':
            context['strategy'] = 'image_group'
            context['needs_spacing'] = False
        elif next_element and next_element.element_type == 'image':
            context['strategy'] = 'image_group_start'
            context['needs_spacing'] = True
        else:
            context['strategy'] = 'standalone'
            context['needs_spacing'] = True
        
        return context
    
    def _analyze_paragraph_context(self, element_index: int) -> Dict[str, Any]:
        """Анализирует контекст параграфа для определения необходимых настроек"""
        context = {
            'needs_spacing_before': False,
            'needs_spacing_after': False,
            'is_title': False,
            'is_quote': False,
            'is_list_item': False
        }
        
        if element_index < len(self.elements):
            current_element = self.elements[element_index]
            
            # Проверяем стиль параграфа
            if current_element.style:
                style_name = current_element.style.lower()
                if 'heading' in style_name or 'title' in style_name:
                    context['is_title'] = True
                    context['needs_spacing_before'] = True
                    context['needs_spacing_after'] = True
                elif 'quote' in style_name:
                    context['is_quote'] = True
                    context['needs_spacing_before'] = True
        
        # Анализируем соседние элементы
        prev_element = self.elements[element_index - 1] if element_index > 0 else None
        if prev_element and prev_element.element_type == 'image':
            context['needs_spacing_before'] = True
        
        return context
    
    def _get_image_temp_path(self, image_element: ImageElement) -> Optional[str]:
        """Получает путь к временному файлу изображения"""
        if not self.improved_image_processor or not self.improved_image_processor.temp_dir:
            return None
            
        temp_path = os.path.join(
            self.improved_image_processor.temp_dir, 
            f"{image_element.image_id}.{image_element.image_format}"
        )
        
        if not os.path.exists(temp_path):
            print(f"⚠️  Файл изображения не найден: {temp_path}")
            return None
            
        return temp_path
    
    def _insert_image_into_new_paragraph(self, document: Document, image_element: ImageElement, temp_path: str, context: Dict[str, Any]) -> bool:
        """Вставляет изображение в новый параграф с учетом контекста"""
        try:
            # Создаем параграф для изображения
            paragraph = document.add_paragraph()
            
            # Применяем настройки на основе контекста
            if context.get('alignment'):
                paragraph.alignment = context['alignment']
            
            # Добавляем изображение
            success = self._add_image_to_paragraph(paragraph, image_element, temp_path)
            
            return success
            
        except Exception as e:
            print(f"❌ Ошибка создания нового параграфа для изображения: {e}")
            return False
    
    def _insert_image_into_existing_paragraph(self, paragraph: Paragraph, image_element: ImageElement, temp_path: str) -> bool:
        """Вставляет изображение в существующий параграф"""
        try:
            return self._add_image_to_paragraph(paragraph, image_element, temp_path)
        except Exception as e:
            print(f"❌ Ошибка вставки в существующий параграф: {e}")
            return False
    
    def _add_image_to_paragraph(self, paragraph: Paragraph, image_element: ImageElement, temp_path: str) -> bool:
        """Добавляет изображение в параграф с правильными размерами"""
        try:
            run = paragraph.add_run()
            
            # Определяем размеры с улучшенной логикой
            width, height = self._calculate_optimal_image_size(image_element)
            
            # Вставляем изображение
            if height:
                run.add_picture(temp_path, width=width, height=height)
            else:
                run.add_picture(temp_path, width=width)
            
            # Устанавливаем выравнивание по центру для изображений
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return True
            
        except Exception as e:
            print(f"❌ Ошибка добавления изображения в параграф: {e}")
            return False
    
    def _calculate_optimal_image_size(self, image_element: ImageElement) -> Tuple[Inches, Optional[Inches]]:
        """Вычисляет оптимальные размеры изображения для документа"""
        # Максимальные размеры
        max_width = 6.0
        max_height = 8.0
        
        if image_element.width and image_element.height:
            # Конвертируем из пикселей в дюймы если нужно
            if image_element.width > 100:  # Похоже на пиксели
                width_inches = image_element.width / 96.0
                height_inches = image_element.height / 96.0
            else:  # Уже в дюймах
                width_inches = image_element.width
                height_inches = image_element.height
                
            # Масштабируем с сохранением пропорций
            if width_inches > max_width:
                scale = max_width / width_inches
                width_inches = max_width
                height_inches = height_inches * scale
                
            if height_inches > max_height:
                scale = max_height / height_inches
                height_inches = max_height
                width_inches = width_inches * scale
                
            return Inches(width_inches), Inches(height_inches)
        else:
            # Размер по умолчанию
            return Inches(4.0), None
    
    def _add_translated_table(self, document: Document, translated_content: str, formatting: Dict[str, Any]):
        """Добавляет переведенную таблицу в документ"""
        try:
            # Разбираем переведенное содержимое таблицы
            rows_content = translated_content.split('\n')
            
            if not rows_content:
                return
            
            # Определяем количество колонок из первой строки
            first_row = rows_content[0].split(' | ')
            cols_count = len(first_row)
            rows_count = len(rows_content)
            
            # Создаем таблицу
            table = document.add_table(rows=rows_count, cols=cols_count)
            
            # Заполняем таблицу
            for row_idx, row_content in enumerate(rows_content):
                cells_content = row_content.split(' | ')
                for col_idx, cell_content in enumerate(cells_content):
                    if col_idx < cols_count:
                        table.cell(row_idx, col_idx).text = cell_content.strip()
            
            # Применяем стиль если есть
            if formatting and 'style' in formatting and formatting['style']:
                try:
                    table.style = formatting['style']
                except:
                    pass  # Игнорируем ошибки стилей
                    
        except Exception as e:
            print(f"Ошибка добавления таблицы: {e}")
    
    def cleanup_temp_files(self):
        """Очищает временные файлы изображений"""
        if self.improved_image_processor:
            self.improved_image_processor.cleanup_temp_files()
    
    def save_as_xml(self, output_path: str) -> bool:
        """
        Сохраняет текст документа в XML формате
        
        Args:
            output_path: Путь для сохранения XML файла
            
        Returns:
            True если сохранение успешно, False иначе
        """
        if not self.elements:
            return False
        
        try:
            # Создаем root элемент
            root = ET.Element("document")
            
            for element in self.elements:
                elem = ET.SubElement(root, element.element_type)
                elem.set("index", str(element.index))
                
                if element.style:
                    elem.set("style", element.style)
                
                # Добавляем текст
                elem.text = element.content
                
                # Добавляем форматирование если есть
                if element.formatting:
                    formatting_elem = ET.SubElement(elem, "formatting")
                    for key, value in element.formatting.items():
                        if value is not None:
                            formatting_elem.set(key, str(value))
            
            # Создаем дерево и сохраняем
            tree = ET.ElementTree(root)
            ET.indent(tree, space="  ", level=0)
            
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            tree.write(output_path, encoding='utf-8', xml_declaration=True)
            return True
            
        except Exception as e:
            print(f"Ошибка сохранения XML: {e}")
            return False
    
    def get_document_statistics(self) -> Dict[str, Any]:
        """Возвращает статистику документа"""
        if not self.elements:
            return {}
        
        total_chars = sum(len(elem.content) for elem in self.elements if elem.element_type != 'image')
        paragraphs = sum(1 for elem in self.elements if elem.element_type == 'paragraph')
        tables = sum(1 for elem in self.elements if elem.element_type == 'table')
        images = sum(1 for elem in self.elements if elem.element_type == 'image')
        
        # Добавляем статистику изображений (используем улучшенный процессор если доступен)
        if self.improved_image_processor and self.improved_image_processor.images:
            image_stats = self.improved_image_processor.get_image_statistics()
        else:
            image_stats = self.improved_image_processor.get_image_statistics() if self.images else {}
        
        stats = {
            'total_elements': len(self.elements),
            'total_characters': total_chars,
            'paragraphs': paragraphs,
            'tables': tables,
            'images': images,
            'average_element_size': total_chars / (len(self.elements) - images) if (len(self.elements) - images) > 0 else 0
        }
        
        # Добавляем детальную статистику изображений
        if image_stats:
            stats.update(image_stats)
        
        return stats
    
    def get_formatting_statistics(self) -> Dict[str, Any]:
        """Возвращает статистику форматирования документа"""
        if not self.elements:
            return {'formatting_complexity': 'none'}
        
        # Собираем данные форматирования всех элементов
        all_formatting_data = []
        for element in self.elements:
            if element.element_type == 'paragraph' and element.formatting:
                all_formatting_data.append(element.formatting)
        
        # Создаем сводку форматирования
        formatting_summary = self.formatting_processor.create_formatting_summary(all_formatting_data)
        
        return formatting_summary
    
    def get_all_text(self) -> str:
        """
        ИСПРАВЛЕНО: Возвращает весь текст документа, сохраняя пустые строки
        в виде специального маркера для точной верстки.
        """
        if not self.elements:
            return ""
        
        EMPTY_PARA_MARKER = "[[EMPTY_PARAGRAPH_MARKER]]"
        
        content_parts = []
        for elem in self.elements:
            if elem.element_type == 'paragraph':
                if not elem.content.strip():
                    content_parts.append(EMPTY_PARA_MARKER)
                else:
                    content_parts.append(elem.content)
            elif elem.element_type == 'table':
                content_parts.append(elem.content)
        
        return '\n\n'.join(content_parts)

    def _log_image_processing_stage(self, stage: str, details: Dict[str, Any]):
        """
        ДЕТАЛЬНОЕ логирование этапов обработки изображений
        
        Args:
            stage: Название этапа (extraction, validation, positioning, insertion)
            details: Детали этапа для логирования
        """
        timestamp = __import__('datetime').datetime.now().strftime("%H:%M:%S.%f")[:-3]
        
        print(f"\n🔍 [{timestamp}] ЭТАП: {stage.upper()}")
        print("=" * 60)
        
        if stage == 'extraction':
            print(f"📥 ИЗВЛЕЧЕНИЕ ИЗОБРАЖЕНИЙ:")
            print(f"  • Путь к файлу: {details.get('file_path', 'Не указан')}")
            print(f"  • Всего параграфов в документе: {details.get('total_paragraphs', 0)}")
            print(f"  • Найдено изображений: {details.get('images_found', 0)}")
            print(f"  • Relationships найдено: {details.get('relationships_count', 0)}")
            print(f"  • Позиций в XML: {details.get('xml_positions_count', 0)}")
            
        elif stage == 'validation':
            print(f"✅ ВАЛИДАЦИЯ ПОЗИЦИЙ:")
            print(f"  • Изображений для проверки: {details.get('total_images', 0)}")
            print(f"  • Валидных позиций: {details.get('valid_positions', 0)}")
            print(f"  • Невалидных позиций: {details.get('invalid_positions', 0)}")
            print(f"  • Исправленных позиций: {details.get('corrected_positions', 0)}")
            print(f"  • Распределенных позиций: {details.get('distributed_positions', 0)}")
            print(f"  • Позиций в конце: {details.get('end_positions', 0)}")
            
        elif stage == 'positioning':
            print(f"📍 ПОЗИЦИОНИРОВАНИЕ:")
            print(f"  • Всего элементов документа: {details.get('total_elements', 0)}")
            print(f"  • Изображений с позициями: {details.get('positioned_images', 0)}")
            print(f"  • Изображений без позиций: {details.get('unpositioned_images', 0)}")
            print(f"  • Параграфов с текстом: {details.get('text_paragraphs', 0)}")
            print(f"  • Таблиц: {details.get('tables_count', 0)}")
            
        elif stage == 'insertion':
            print(f"🔄 ВСТАВКА В ПЕРЕВЕДЕННЫЙ ДОКУМЕНТ:")
            print(f"  • Всего элементов для обработки: {details.get('total_elements', 0)}")
            print(f"  • Успешно вставлено изображений: {details.get('images_inserted', 0)}")
            print(f"  • Ошибок вставки изображений: {details.get('images_failed', 0)}")
            print(f"  • Обработано параграфов: {details.get('paragraphs_processed', 0)}")
            print(f"  • Обработано таблиц: {details.get('tables_processed', 0)}")
            
        elif stage == 'debug_analysis':
            print(f"🐛 ДИАГНОСТИЧЕСКИЙ АНАЛИЗ:")
            if 'image_positions_map' in details:
                print(f"  • Карта позиций изображений:")
                for img_id, pos in details['image_positions_map'].items():
                    print(f"    - {img_id}: позиция {pos}")
            
            if 'paragraph_analysis' in details:
                print(f"  • Анализ параграфов:")
                for i, para_info in enumerate(details['paragraph_analysis'][:10]):  # Первые 10
                    print(f"    - Параграф {i}: {para_info}")
                if len(details['paragraph_analysis']) > 10:
                    print(f"    ... и еще {len(details['paragraph_analysis']) - 10} параграфов")
        
        print("=" * 60)
        print()
    
    def _determine_smart_distribution_strategy(self, images: List[ImageElement], text_paragraphs: int, current_elements: int) -> str:
        """
        Определяет оптимальную стратегию распределения изображений без позиций
        
        Args:
            images: Изображения без позиций
            text_paragraphs: Количество текстовых параграфов
            current_elements: Текущее количество элементов
            
        Returns:
            Название стратегии
        """
        images_count = len(images)
        
        # Анализ соотношений для выбора стратегии
        if images_count <= 2:
            return 'end_placement'  # Мало изображений - в конец
        elif images_count <= text_paragraphs // 4:
            return 'strategic_insertion'  # Стратегическая вставка в ключевые места
        elif images_count <= text_paragraphs // 2:
            return 'proportional_distribution'  # Пропорциональное распределение
        elif text_paragraphs > 20:
            return 'chapter_clustering'  # Группировка по разделам
        else:
            return 'end_placement'  # Слишком много изображений - в конец
    
    def _insert_images_strategically(self, elements: List[DocumentElement], images: List[ImageElement], start_index: int) -> List[DocumentElement]:
        """Стратегическая вставка изображений в ключевые места документа"""
        print(f"🎯 СТРАТЕГИЧЕСКАЯ ВСТАВКА: {len(images)} изображений")
        
        # Находим стратегические позиции (начало разделов, после заголовков, etc.)
        strategic_positions = self._find_strategic_positions(elements)
        
        current_index = start_index
        images_inserted = 0
        
        for i, image in enumerate(images):
            if i < len(strategic_positions):
                # Вставляем в стратегическую позицию
                position = strategic_positions[i]
                image_element = DocumentElement(
                    element_type='image',
                    content=f"[IMAGE: {image.image_id}]",
                    original_element=None,
                    index=current_index,
                    image_element=image
                )
                elements.insert(position + i, image_element)
                current_index += 1
                images_inserted += 1
                print(f"🎯 СТРАТЕГИЧЕСКИ: Изображение {image.image_id} вставлено в позицию {position + i}")
            else:
                # Остальные в конец
                image_element = DocumentElement(
                    element_type='image',
                    content=f"[IMAGE: {image.image_id}]",
                    original_element=None,
                    index=current_index,
                    image_element=image
                )
                elements.append(image_element)
                current_index += 1
                print(f"📌 ДОПОЛНИТЕЛЬНО: Изображение {image.image_id} добавлено в конец")
        
        print(f"🎯 РЕЗУЛЬТАТ: {images_inserted} изображений размещено стратегически")
        return elements
    
    def _distribute_images_proportionally(self, elements: List[DocumentElement], images: List[ImageElement], start_index: int) -> List[DocumentElement]:
        """Пропорциональное распределение изображений по документу"""
        print(f"📊 ПРОПОРЦИОНАЛЬНОЕ РАСПРЕДЕЛЕНИЕ: {len(images)} изображений")
        
        text_elements = [i for i, elem in enumerate(elements) if elem.element_type == 'paragraph']
        
        if not text_elements:
            return self._place_images_at_end(elements, images, start_index)
        
        # Вычисляем позиции для равномерного распределения
        step = len(text_elements) // (len(images) + 1) if len(images) > 0 else 1
        if step < 1:
            step = 1
        
        current_index = start_index
        
        for i, image in enumerate(images):
            if i * step < len(text_elements):
                # Вставляем после соответствующего текстового элемента
                position = text_elements[min(i * step, len(text_elements) - 1)] + 1 + i
                image_element = DocumentElement(
                    element_type='image',
                    content=f"[IMAGE: {image.image_id}]",
                    original_element=None,
                    index=current_index,
                    image_element=image
                )
                if position < len(elements):
                    elements.insert(position, image_element)
                else:
                    elements.append(image_element)
                current_index += 1
                print(f"📊 ПРОПОРЦИОНАЛЬНО: Изображение {image.image_id} вставлено в позицию {position}")
            else:
                # Остальные в конец
                image_element = DocumentElement(
                    element_type='image',
                    content=f"[IMAGE: {image.image_id}]",
                    original_element=None,
                    index=current_index,
                    image_element=image
                )
                elements.append(image_element)
                current_index += 1
                print(f"📌 ДОПОЛНИТЕЛЬНО: Изображение {image.image_id} добавлено в конец")
        
        return elements
    
    def _cluster_images_by_chapters(self, elements: List[DocumentElement], images: List[ImageElement], start_index: int) -> List[DocumentElement]:
        """Группировка изображений по разделам документа"""
        print(f"📚 ГРУППИРОВКА ПО РАЗДЕЛАМ: {len(images)} изображений")
        
        # Находим предполагаемые разделы (заголовки, значительные отступы в тексте)
        chapter_positions = self._find_chapter_boundaries(elements)
        
        if not chapter_positions:
            return self._distribute_images_proportionally(elements, images, start_index)
        
        # Распределяем изображения по разделам
        images_per_chapter = len(images) // len(chapter_positions)
        remaining_images = len(images) % len(chapter_positions)
        
        current_index = start_index
        image_idx = 0
        
        for chapter_idx, chapter_pos in enumerate(chapter_positions):
            # Количество изображений для этого раздела
            images_for_chapter = images_per_chapter + (1 if chapter_idx < remaining_images else 0)
            
            for i in range(images_for_chapter):
                if image_idx < len(images):
                    image = images[image_idx]
                    image_element = DocumentElement(
                        element_type='image',
                        content=f"[IMAGE: {image.image_id}]",
                        original_element=None,
                        index=current_index,
                        image_element=image
                    )
                    # Вставляем после начала раздела
                    insert_position = chapter_pos + 1 + i + sum(images_per_chapter + (1 if j < remaining_images else 0) for j in range(chapter_idx))
                    if insert_position < len(elements):
                        elements.insert(insert_position, image_element)
                    else:
                        elements.append(image_element)
                    
                    current_index += 1
                    image_idx += 1
                    print(f"📚 РАЗДЕЛ {chapter_idx + 1}: Изображение {image.image_id} добавлено")
        
        # Остальные изображения в конец
        while image_idx < len(images):
            image = images[image_idx]
            image_element = DocumentElement(
                element_type='image',
                content=f"[IMAGE: {image.image_id}]",
                original_element=None,
                index=current_index,
                image_element=image
            )
            elements.append(image_element)
            current_index += 1
            image_idx += 1
            print(f"📌 ДОПОЛНИТЕЛЬНО: Изображение {image.image_id} добавлено в конец")
        
        return elements
    
    def _place_images_at_end(self, elements: List[DocumentElement], images: List[ImageElement], start_index: int) -> List[DocumentElement]:
        """Размещение изображений в конце документа (исходная стратегия)"""
        print(f"📌 РАЗМЕЩЕНИЕ В КОНЦЕ: {len(images)} изображений")
        
        current_index = start_index
        
        for image in images:
            image_element = DocumentElement(
                element_type='image',
                content=f"[IMAGE: {image.image_id}]",
                original_element=None,
                index=current_index,
                image_element=image
            )
            elements.append(image_element)
            current_index += 1
            print(f"📌 Изображение {image.image_id} добавлено в конец документа")
        
        return elements
    
    def _find_strategic_positions(self, elements: List[DocumentElement]) -> List[int]:
        """Находит стратегические позиции для вставки изображений"""
        positions = []
        
        # Ищем после каждого 3-4 параграфа
        paragraph_count = 0
        for i, element in enumerate(elements):
            if element.element_type == 'paragraph':
                paragraph_count += 1
                if paragraph_count % 3 == 0:  # Каждый третий параграф
                    positions.append(i)
        
        return positions[:10]  # Ограничиваем количество позиций
    
    def _find_chapter_boundaries(self, elements: List[DocumentElement]) -> List[int]:
        """Находит границы разделов в документе"""
        boundaries = []
        
        # Простая эвристика: каждые 10-15 параграфов
        paragraph_count = 0
        for i, element in enumerate(elements):
            if element.element_type == 'paragraph':
                paragraph_count += 1
                if paragraph_count % 12 == 0:  # Каждые 12 параграфов
                    boundaries.append(i)
        
        return boundaries[:5]  # Ограничиваем количество разделов

    def coordinate_image_processing_components(self) -> Dict[str, Any]:
        """
        КООРДИНАЦИЯ всех компонентов обработки изображений
        Обеспечивает синхронизацию между ImageAdapter, ImprovedImageProcessor и DocumentProcessor
        
        Returns:
            Отчет о координации компонентов
        """
        coordination_report = {
            'timestamp': __import__('datetime').datetime.now().isoformat(),
            'components_status': {},
            'synchronization_issues': [],
            'performance_metrics': {},
            'recommendations': []
        }
        
        print(f"🔄 КООРДИНАЦИЯ КОМПОНЕНТОВ: Синхронизация системы обработки изображений")
        print("=" * 70)
        
        # === ЭТАП 1: ПРОВЕРКА СОСТОЯНИЯ КОМПОНЕНТОВ ===
        try:
            # Проверяем DocumentProcessor
            doc_status = self._check_document_processor_status()
            coordination_report['components_status']['DocumentProcessor'] = doc_status
            
            # Проверяем ImprovedImageProcessor
            image_proc_status = self._check_improved_image_processor_status()
            coordination_report['components_status']['ImprovedImageProcessor'] = image_proc_status
            
            # Проверяем ImageAdapter
            adapter_status = self._check_image_adapter_status()
            coordination_report['components_status']['ImageAdapter'] = adapter_status
            
            print(f"📊 СОСТОЯНИЕ КОМПОНЕНТОВ:")
            print(f"  • DocumentProcessor: {doc_status['status']}")
            print(f"  • ImprovedImageProcessor: {image_proc_status['status']}")
            print(f"  • ImageAdapter: {adapter_status['status']}")
            
        except Exception as e:
            coordination_report['synchronization_issues'].append(f"Ошибка проверки компонентов: {e}")
            print(f"❌ Ошибка проверки компонентов: {e}")
        
        # === ЭТАП 2: СИНХРОНИЗАЦИЯ ДАННЫХ ===
        try:
            sync_result = self._synchronize_component_data()
            coordination_report['synchronization_result'] = sync_result
            
            if sync_result['success']:
                print(f"✅ СИНХРОНИЗАЦИЯ: Данные успешно синхронизированы")
                print(f"  • Изображений синхронизировано: {sync_result['images_synchronized']}")
                print(f"  • Позиций исправлено: {sync_result['positions_corrected']}")
            else:
                print(f"❌ СИНХРОНИЗАЦИЯ: Обнаружены проблемы")
                coordination_report['synchronization_issues'].extend(sync_result['issues'])
                
        except Exception as e:
            coordination_report['synchronization_issues'].append(f"Ошибка синхронизации: {e}")
            print(f"❌ Ошибка синхронизации: {e}")
        
        # === ЭТАП 3: АНАЛИЗ ПРОИЗВОДИТЕЛЬНОСТИ ===
        try:
            performance_metrics = self._analyze_performance_metrics()
            coordination_report['performance_metrics'] = performance_metrics
            
            print(f"📈 МЕТРИКИ ПРОИЗВОДИТЕЛЬНОСТИ:")
            print(f"  • Время извлечения изображений: {performance_metrics.get('extraction_time', 'N/A')}")
            print(f"  • Время валидации: {performance_metrics.get('validation_time', 'N/A')}")
            print(f"  • Время вставки: {performance_metrics.get('insertion_time', 'N/A')}")
            print(f"  • Успешность позиционирования: {performance_metrics.get('positioning_success_rate', 'N/A')}%")
            
        except Exception as e:
            coordination_report['synchronization_issues'].append(f"Ошибка анализа производительности: {e}")
            print(f"❌ Ошибка анализа производительности: {e}")
        
        # === ЭТАП 4: РЕКОМЕНДАЦИИ ===
        recommendations = self._generate_coordination_recommendations(coordination_report)
        coordination_report['recommendations'] = recommendations
        
        print(f"💡 РЕКОМЕНДАЦИИ:")
        for i, rec in enumerate(recommendations, 1):
            print(f"  {i}. {rec}")
        
        print("=" * 70)
        
        return coordination_report
    
    def _check_document_processor_status(self) -> Dict[str, Any]:
        """Проверяет состояние DocumentProcessor"""
        status = {
            'status': 'healthy',
            'document_loaded': self.document is not None,
            'elements_count': len(self.elements),
            'images_count': len(self.images),
            'file_path': self.file_path,
            'issues': []
        }
        
        # Проверяем основные компоненты
        if not self.document:
            status['issues'].append("Документ не загружен")
            status['status'] = 'warning'
        
        if not self.elements:
            status['issues'].append("Элементы не извлечены")
            status['status'] = 'warning'
        
        if not self.file_path:
            status['issues'].append("Путь к файлу не установлен")
            status['status'] = 'warning'
        
        return status
    
    def _check_improved_image_processor_status(self) -> Dict[str, Any]:
        """Проверяет состояние ImprovedImageProcessor"""
        status = {
            'status': 'healthy',
            'temp_dir_exists': False,
            'images_extracted': 0,
            'last_extraction_time': None,
            'issues': []
        }
        
        if self.improved_image_processor:
            status['temp_dir_exists'] = bool(self.improved_image_processor.temp_dir and 
                                           os.path.exists(self.improved_image_processor.temp_dir))
            status['images_extracted'] = len(getattr(self.improved_image_processor, 'images', []))
            
            if not status['temp_dir_exists']:
                status['issues'].append("Временная папка не создана")
                status['status'] = 'warning'
        else:
            status['issues'].append("ImprovedImageProcessor не инициализирован")
            status['status'] = 'error'
        
        return status
    
    def _check_image_adapter_status(self) -> Dict[str, Any]:
        """Проверяет состояние ImageAdapter"""
        status = {
            'status': 'healthy',
            'conversion_available': True,
            'issues': []
        }
        
        # Проверяем доступность методов ImageAdapter
        try:
            from image_adapter import ImageAdapter
            if not hasattr(ImageAdapter, 'convert_to_image_element'):
                status['issues'].append("Метод convert_to_image_element недоступен")
                status['status'] = 'error'
            if not hasattr(ImageAdapter, 'convert_list_to_image_elements'):
                status['issues'].append("Метод convert_list_to_image_elements недоступен")
                status['status'] = 'error'
        except ImportError:
            status['issues'].append("ImageAdapter не может быть импортирован")
            status['status'] = 'error'
        
        return status
    
    def _synchronize_component_data(self) -> Dict[str, Any]:
        """Синхронизирует данные между компонентами"""
        sync_result = {
            'success': True,
            'images_synchronized': 0,
            'positions_corrected': 0,
            'issues': []
        }
        
        try:
            # Синхронизируем данные изображений
            if self.improved_image_processor and hasattr(self.improved_image_processor, 'images'):
                processor_images = self.improved_image_processor.images
                adapter_images = self.images
                
                # Сравниваем количество изображений
                if len(processor_images) != len(adapter_images):
                    sync_result['issues'].append(f"Несоответствие количества изображений: processor={len(processor_images)}, adapter={len(adapter_images)}")
                    sync_result['success'] = False
                
                # Проверяем соответствие позиций
                positions_corrected = 0
                for proc_img, adapt_img in zip(processor_images, adapter_images):
                    if proc_img.paragraph_index != adapt_img.paragraph_index:
                        print(f"🔄 КОРРЕКЦИЯ: Синхронизация позиции для {adapt_img.image_id}: {adapt_img.paragraph_index} -> {proc_img.paragraph_index}")
                        adapt_img.paragraph_index = proc_img.paragraph_index
                        positions_corrected += 1
                
                sync_result['images_synchronized'] = len(adapter_images)
                sync_result['positions_corrected'] = positions_corrected
                
        except Exception as e:
            sync_result['issues'].append(f"Ошибка синхронизации данных: {e}")
            sync_result['success'] = False
        
        return sync_result
    
    def _analyze_performance_metrics(self) -> Dict[str, Any]:
        """Анализирует метрики производительности системы"""
        metrics = {
            'extraction_time': 'N/A',
            'validation_time': 'N/A',
            'insertion_time': 'N/A',
            'positioning_success_rate': 0,
            'memory_usage': 'N/A'
        }
        
        try:
            # Вычисляем успешность позиционирования
            if self.images:
                positioned_count = len([img for img in self.images if img.paragraph_index is not None])
                metrics['positioning_success_rate'] = round((positioned_count / len(self.images)) * 100, 1)
            
            # Анализируем использование памяти
            import psutil
            process = psutil.Process()
            memory_info = process.memory_info()
            metrics['memory_usage'] = f"{memory_info.rss / 1024 / 1024:.1f} MB"
            
        except Exception as e:
            metrics['analysis_error'] = str(e)
        
        return metrics
    
    def _generate_coordination_recommendations(self, report: Dict[str, Any]) -> List[str]:
        """Генерирует рекомендации по улучшению координации компонентов"""
        recommendations = []
        
        # Анализируем проблемы компонентов
        if report['components_status'].get('DocumentProcessor', {}).get('status') != 'healthy':
            recommendations.append("Проверить состояние DocumentProcessor и загрузку документа")
        
        if report['components_status'].get('ImprovedImageProcessor', {}).get('status') != 'healthy':
            recommendations.append("Переинициализировать ImprovedImageProcessor и создать временную папку")
        
        if report['components_status'].get('ImageAdapter', {}).get('status') != 'healthy':
            recommendations.append("Проверить доступность ImageAdapter и его методов")
        
        # Анализируем синхронизацию
        if report.get('synchronization_result', {}).get('success') == False:
            recommendations.append("Выполнить ресинхронизацию данных между компонентами")
        
        # Анализируем производительность
        success_rate = report.get('performance_metrics', {}).get('positioning_success_rate', 0)
        if success_rate < 80:
            recommendations.append(f"Улучшить алгоритм позиционирования (текущая успешность: {success_rate}%)")
        
        if not recommendations:
            recommendations.append("Система работает стабильно, дополнительные действия не требуются")
        
        return recommendations 

    def _track_image_positions(self, stage: str, images: List[ImageElement], additional_info: Dict[str, Any] = None):
        """
        ОТСЛЕЖИВАНИЕ позиций изображений на каждом этапе обработки
        
        Args:
            stage: Этап обработки (extraction, validation, positioning, insertion)
            images: Список изображений для отслеживания
            additional_info: Дополнительная информация об этапе
        """
        if not self.position_tracker['tracking_enabled']:
            return
            
        timestamp = __import__('datetime').datetime.now().isoformat()
        
        # Сохраняем позиции изображений
        stage_positions = {}
        for image in images:
            stage_positions[image.image_id] = {
                'paragraph_index': image.paragraph_index,
                'image_format': image.image_format,
                'width': image.width,
                'height': image.height,
                'timestamp': timestamp
            }
        
        self.position_tracker[f'{stage}_stage'] = stage_positions
        
        # Записываем в историю изменений
        history_entry = {
            'timestamp': timestamp,
            'stage': stage,
            'images_count': len(images),
            'positioned_count': len([img for img in images if img.paragraph_index is not None]),
            'unpositioned_count': len([img for img in images if img.paragraph_index is None]),
            'additional_info': additional_info or {}
        }
        
        self.position_tracker['position_history'].append(history_entry)
        
        # Логируем изменения позиций если это не первый этап
        if stage != 'extraction':
            self._log_position_changes(stage, stage_positions)
    
    def _log_position_changes(self, current_stage: str, current_positions: Dict[str, Any]):
        """Логирует изменения позиций между этапами"""
        previous_stage_map = {
            'validation': 'extraction',
            'positioning': 'validation', 
            'insertion': 'positioning'
        }
        
        previous_stage = previous_stage_map.get(current_stage)
        if not previous_stage:
            return
            
        previous_positions = self.position_tracker.get(f'{previous_stage}_stage', {})
        
        changes_detected = 0
        for image_id, current_info in current_positions.items():
            if image_id in previous_positions:
                prev_pos = previous_positions[image_id]['paragraph_index']
                curr_pos = current_info['paragraph_index']
                
                if prev_pos != curr_pos:
                    changes_detected += 1
                    print(f"🔄 ТРЕКИНГ: {image_id} позиция изменилась на этапе {current_stage}: {prev_pos} -> {curr_pos}")
        
        if changes_detected > 0:
            print(f"📊 ТРЕКИНГ: На этапе {current_stage} изменено позиций: {changes_detected}")
    
    def get_position_tracking_report(self) -> Dict[str, Any]:
        """
        Возвращает детальный отчет по отслеживанию позиций изображений
        
        Returns:
            Отчет о треке позиций на всех этапах
        """
        report = {
            'tracking_enabled': self.position_tracker['tracking_enabled'],
            'stages_tracked': [],
            'position_stability': {},
            'problematic_images': [],
            'summary': {}
        }
        
        # Анализируем каждый этап
        stages = ['extraction', 'validation', 'positioning', 'insertion']
        for stage in stages:
            stage_key = f'{stage}_stage'
            if stage_key in self.position_tracker and self.position_tracker[stage_key]:
                report['stages_tracked'].append(stage)
        
        # Анализируем стабильность позиций
        if len(report['stages_tracked']) > 1:
            report['position_stability'] = self._analyze_position_stability()
        
        # Находим проблемные изображения
        report['problematic_images'] = self._identify_problematic_images()
        
        # Создаем сводку
        if self.position_tracker['position_history']:
            latest_entry = self.position_tracker['position_history'][-1]
            report['summary'] = {
                'total_images': latest_entry['images_count'],
                'positioned_images': latest_entry['positioned_count'],
                'unpositioned_images': latest_entry['unpositioned_count'],
                'positioning_success_rate': round((latest_entry['positioned_count'] / latest_entry['images_count']) * 100, 1) if latest_entry['images_count'] > 0 else 0,
                'stages_completed': len(report['stages_tracked']),
                'last_update': latest_entry['timestamp']
            }
        
        return report
    
    def _analyze_position_stability(self) -> Dict[str, Any]:
        """Анализирует стабильность позиций изображений между этапами"""
        stability = {
            'stable_images': 0,
            'unstable_images': 0,
            'stability_rate': 0,
            'stage_transitions': {}
        }
        
        stages = ['extraction', 'validation', 'positioning', 'insertion']
        
        # Получаем все уникальные ID изображений
        all_image_ids = set()
        for stage in stages:
            stage_positions = self.position_tracker.get(f'{stage}_stage', {})
            all_image_ids.update(stage_positions.keys())
        
        # Анализируем каждое изображение
        for image_id in all_image_ids:
            positions_across_stages = []
            for stage in stages:
                stage_positions = self.position_tracker.get(f'{stage}_stage', {})
                if image_id in stage_positions:
                    positions_across_stages.append(stage_positions[image_id]['paragraph_index'])
            
            # Проверяем стабильность
            if len(set(positions_across_stages)) == 1:
                stability['stable_images'] += 1
            else:
                stability['unstable_images'] += 1
        
        # Вычисляем процент стабильности
        total_images = stability['stable_images'] + stability['unstable_images']
        if total_images > 0:
            stability['stability_rate'] = round((stability['stable_images'] / total_images) * 100, 1)
        
        return stability
    
    def _identify_problematic_images(self) -> List[Dict[str, Any]]:
        """Выявляет изображения с проблемами позиционирования"""
        problematic = []
        
        stages = ['extraction', 'validation', 'positioning', 'insertion']
        
        # Получаем все уникальные ID изображений
        all_image_ids = set()
        for stage in stages:
            stage_positions = self.position_tracker.get(f'{stage}_stage', {})
            all_image_ids.update(stage_positions.keys())
        
        for image_id in all_image_ids:
            issues = []
            position_history = []
            
            for stage in stages:
                stage_positions = self.position_tracker.get(f'{stage}_stage', {})
                if image_id in stage_positions:
                    pos = stage_positions[image_id]['paragraph_index']
                    position_history.append({'stage': stage, 'position': pos})
                    
                    # Проверяем на проблемы
                    if pos is None:
                        issues.append(f"Отсутствует позиция на этапе {stage}")
                    elif isinstance(pos, int) and pos < 0:
                        issues.append(f"Отрицательная позиция на этапе {stage}: {pos}")
            
            # Проверяем на частые изменения позиций
            positions = [p['position'] for p in position_history if p['position'] is not None]
            if len(set(positions)) > 2:
                issues.append(f"Частые изменения позиций: {positions}")
            
            if issues:
                problematic.append({
                    'image_id': image_id,
                    'issues': issues,
                    'position_history': position_history
                })
        
        return problematic
    
    def enable_position_tracking(self):
        """Включает отслеживание позиций"""
        self.position_tracker['tracking_enabled'] = True
        print("✅ Отслеживание позиций изображений включено")
    
    def disable_position_tracking(self):
        """Отключает отслеживание позиций"""
        self.position_tracker['tracking_enabled'] = False
        print("❌ Отслеживание позиций изображений отключено")
    
    def clear_position_tracking_history(self):
        """Очищает историю отслеживания позиций"""
        self.position_tracker = {
            'extraction_stage': {},
            'validation_stage': {},
            'positioning_stage': {},
            'insertion_stage': {},
            'position_history': [],
            'tracking_enabled': self.position_tracker['tracking_enabled']
        }
        print("🗑️  История отслеживания позиций очищена")

    def run_comprehensive_image_positioning_test(self, test_document_path: str = None) -> Dict[str, Any]:
        """
        КОМПЛЕКСНОЕ ТЕСТИРОВАНИЕ исправленной системы позиционирования изображений
        
        Args:
            test_document_path: Путь к тестовому документу (опционально)
            
        Returns:
            Детальный отчет о тестировании
        """
        test_results = {
            'timestamp': __import__('datetime').datetime.now().isoformat(),
            'test_document': test_document_path or self.file_path,
            'system_components': {},
            'position_accuracy': {},
            'performance_metrics': {},
            'issues_found': [],
            'recommendations': [],
            'overall_success': False
        }
        
        print(f"🧪 КОМПЛЕКСНОЕ ТЕСТИРОВАНИЕ СИСТЕМЫ ПОЗИЦИОНИРОВАНИЯ ИЗОБРАЖЕНИЙ")
        print("=" * 80)
        
        try:
            # === ТЕСТ 1: ПРОВЕРКА КОМПОНЕНТОВ ===
            print(f"📋 ТЕСТ 1: Проверка компонентов системы")
            components_report = self.coordinate_image_processing_components()
            test_results['system_components'] = components_report
            
            if any(status.get('status') != 'healthy' for status in components_report['components_status'].values()):
                test_results['issues_found'].append("Обнаружены проблемы в компонентах системы")
            else:
                print("✅ Все компоненты системы работают корректно")
            
            # === ТЕСТ 2: ТЕСТИРОВАНИЕ ПОЗИЦИОНИРОВАНИЯ ===
            print(f"\n📋 ТЕСТ 2: Тестирование точности позиционирования")
            
            if test_document_path and test_document_path != self.file_path:
                # Загружаем тестовый документ
                if self.load_document(test_document_path):
                    print(f"✅ Тестовый документ загружен: {test_document_path}")
                else:
                    test_results['issues_found'].append(f"Не удалось загрузить тестовый документ: {test_document_path}")
                    return test_results
            
            if self.document:
                # Выполняем полный цикл извлечения и позиционирования
                elements = self.extract_text_elements()
                
                # Анализируем результаты
                positioning_report = self.get_position_tracking_report()
                test_results['position_accuracy'] = positioning_report
                
                success_rate = positioning_report.get('summary', {}).get('positioning_success_rate', 0)
                print(f"📊 Успешность позиционирования: {success_rate}%")
                
                if success_rate < 70:
                    test_results['issues_found'].append(f"Низкая успешность позиционирования: {success_rate}%")
                elif success_rate >= 90:
                    print("✅ Отличная точность позиционирования")
                else:
                    print("⚠️  Удовлетворительная точность позиционирования")
            
            # === ТЕСТ 3: ПРОИЗВОДИТЕЛЬНОСТЬ ===
            print(f"\n📋 ТЕСТ 3: Анализ производительности")
            
            import time
            start_time = time.time()
            
            # Создаем тестовый переведенный документ
            mock_translation_results = self._create_mock_translation_results()
            translated_doc = self.create_translated_document(mock_translation_results)
            
            end_time = time.time()
            processing_time = end_time - start_time
            
            performance_metrics = {
                'total_processing_time': round(processing_time, 2),
                'images_processed': len(self.images),
                'elements_processed': len(self.elements),
                'processing_speed': round(len(self.elements) / processing_time, 2) if processing_time > 0 else 0
            }
            
            test_results['performance_metrics'] = performance_metrics
            print(f"📈 Время обработки: {processing_time:.2f} сек")
            print(f"📈 Скорость обработки: {performance_metrics['processing_speed']:.2f} элементов/сек")
            
            if translated_doc:
                print("✅ Создание переведенного документа успешно")
            else:
                test_results['issues_found'].append("Не удалось создать переведенный документ")
            
            # === ТЕСТ 4: СТРЕСС-ТЕСТИРОВАНИЕ ===
            print(f"\n📋 ТЕСТ 4: Стресс-тестирование")
            stress_test_results = self._run_stress_test()
            test_results['stress_test'] = stress_test_results
            
            if stress_test_results['success']:
                print("✅ Стресс-тест пройден успешно")
            else:
                test_results['issues_found'].append("Обнаружены проблемы при стресс-тестировании")
            
            # === ФИНАЛЬНАЯ ОЦЕНКА ===
            test_results['overall_success'] = len(test_results['issues_found']) == 0
            
            if test_results['overall_success']:
                print(f"\n🎉 ТЕСТИРОВАНИЕ ЗАВЕРШЕНО УСПЕШНО!")
                print(f"✅ Система позиционирования изображений работает корректно")
                test_results['recommendations'].append("Система готова к продакшн использованию")
            else:
                print(f"\n⚠️  ТЕСТИРОВАНИЕ ВЫЯВИЛО ПРОБЛЕМЫ:")
                for issue in test_results['issues_found']:
                    print(f"  ❌ {issue}")
                test_results['recommendations'].append("Требуется устранение выявленных проблем")
            
            print("=" * 80)
            
        except Exception as e:
            test_results['issues_found'].append(f"Критическая ошибка тестирования: {e}")
            test_results['overall_success'] = False
            print(f"❌ Критическая ошибка тестирования: {e}")
        
        return test_results
    
    def _create_mock_translation_results(self) -> List[Any]:
        """Создает mock-объекты результатов перевода для тестирования"""
        from types import SimpleNamespace
        
        mock_results = []
        for element in self.elements:
            if element.element_type in ['paragraph', 'table']:
                mock_result = SimpleNamespace()
                mock_result.success = True
                mock_result.translated_text = f"[ТЕСТ] Переведенный текст для элемента {element.index}"
                mock_results.append(mock_result)
        
        return mock_results
    
    def _run_stress_test(self) -> Dict[str, Any]:
        """Выполняет стресс-тестирование системы"""
        stress_results = {
            'success': True,
            'iterations_completed': 0,
            'max_iterations': 5,
            'errors_encountered': [],
            'performance_degradation': False
        }
        
        try:
            initial_time = None
            
            for i in range(stress_results['max_iterations']):
                start_time = time.time()
                
                # Повторно выполняем извлечение и валидацию
                if self.file_path:
                    image_infos = self.improved_image_processor.extract_images_from_docx(self.file_path)
                    test_images = ImageAdapter.convert_list_to_image_elements(image_infos)
                    validated_images = self._validate_and_correct_image_positions(test_images)
                
                end_time = time.time()
                iteration_time = end_time - start_time
                
                if initial_time is None:
                    initial_time = iteration_time
                elif iteration_time > initial_time * 2:  # Если время увеличилось в 2 раза
                    stress_results['performance_degradation'] = True
                
                stress_results['iterations_completed'] += 1
                print(f"  Итерация {i+1}/{stress_results['max_iterations']}: {iteration_time:.2f} сек")
                
        except Exception as e:
            stress_results['errors_encountered'].append(str(e))
            stress_results['success'] = False
        
        return stress_results